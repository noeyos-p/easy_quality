"""
RAG 챗봇 API v14.0 + Agent (OpenAI)

 v14.0 변경사항:
- LLM 백엔드 변경: Z.AI → OpenAI GPT 계열
- 에이전트 시스템 통합 (모든 서브 에이전트 OpenAI 사용)
- LLM as a Judge 평가 시스템 (RDB 검증 포함)
- LangSmith 추적 지원 및 최적화
"""

#  .env 파일 자동 로드 (다른 import보다 먼저!)
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Depends, status, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from pydantic import BaseModel, Field, field_validator
from typing import List, Dict, Optional, Literal
from contextlib import asynccontextmanager
import torch
import time
import re
import uuid
import asyncio
import json
import os
import jwt
from io import BytesIO

from backend.sql_store import SQLStore
sql_store = SQLStore()
# sql_store.init_db()  #  main()으로 이동하여 중복 호출 방지

# RAG 모듈 - 레거시 (폴백용)
# RAG 모듈 - 레거시 (폴백용) 제거됨
# LangGraph 파이프라인이 전적으로 처리

from sentence_transformers import SentenceTransformer
from backend import vector_store
from backend.vector_store import embed_text
# from backend.prompt import build_rag_prompt, build_chunk_prompt (제거됨)
from backend.llm import (
    get_llm_response,
    ZaiLLM,
    OllamaLLM,
    analyze_search_results,
    HUGGINGFACE_MODELS,
)


# ═══════════════════════════════════════════════════════════════════════════
# 설정 및 모델
# ═══════════════════════════════════════════════════════════════════════════

DEFAULT_CHUNK_SIZE = 500
DEFAULT_OVERLAP = 50
DEFAULT_CHUNK_METHOD = "article"
DEFAULT_N_RESULTS = 7
DEFAULT_SIMILARITY_THRESHOLD = 0.30
USE_LANGGRAPH = True
FORCED_LLM_MODEL = "gpt-4o"

def resolve_effective_llm_model(requested_model: Optional[str]) -> str:
    """클라이언트 요청값과 무관하게 서버 모델을 gpt-4o로 고정."""
    return FORCED_LLM_MODEL

class SearchRequest(BaseModel):
    query: str
    collection: str = "documents"
    n_results: int = DEFAULT_N_RESULTS
    model: str = "multilingual-e5-small"
    filter_doc: Optional[str] = None
    similarity_threshold: Optional[float] = None

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    user_id: Optional[int] = None
    collection: str = "documents"
    n_results: int = DEFAULT_N_RESULTS
    embedding_model: str = "multilingual-e5-small"
    llm_model: str = "gpt-4o"
    llm_backend: str = "openai"
    filter_doc: Optional[str] = None
    similarity_threshold: Optional[float] = None

class AskRequest(BaseModel):
    query: str
    collection: str = "documents"
    n_results: int = DEFAULT_N_RESULTS
    embedding_model: str = "multilingual-e5-small"
    llm_model: str = "gpt-4o"
    llm_backend: str = "openai"
    temperature: float = 0.7
    filter_doc: Optional[str] = None
    language: str = "ko"
    max_tokens: int = 512
    similarity_threshold: Optional[float] = None
    include_sources: bool = True

class LLMRequest(BaseModel):
    prompt: str
    model: str = "qwen2.5:3b"
    backend: str = "ollama"
    max_tokens: int = 256
    temperature: float = 0.1

class DeleteDocRequest(BaseModel):
    doc_name: str
    collection: str = "documents"
    delete_from_neo4j: bool = True

class DeleteDocsBatchRequest(BaseModel):
    doc_names: List[str]
    collection: str = "documents"
    delete_from_neo4j: bool = True

class SaveDocRequest(BaseModel):
    doc_name: str
    content: str
    collection: str = "documents"
    model: str = "multilingual-e5-small"

PRESET_MODELS = {
    "multilingual-e5-small": "intfloat/multilingual-e5-small",
}

device = "cuda" if torch.cuda.is_available() else "cpu"
chat_histories: Dict[str, List[Dict]] = {}
chat_results: Dict[str, Dict] = {}
chat_queue: asyncio.Queue = asyncio.Queue()

# ticket은 단조 증가(서버 살아있는 동안 1,2,3,4...)
next_ticket: int = 0
# pending은 request_id만 들고 있지 말고 ticket도 같이 들고 있게
# [{"request_id": "...", "ticket": 1, "kind": "rag"|"agent"} ...]
chat_pending: List[Dict] = []
ticket_lock = asyncio.Lock()

_graph_store = None

QUEUE_STATE_FILE = "backend/queue_state.json"
queue_lock = asyncio.Lock()

# 🔄 백그라운드 작업 상태 관리 (문서 업로드, 수정 등)
processing_tasks: Dict[str, Dict] = {}

def update_task_status(task_id: str, status: str, message: str = "", **kwargs):
    """작업 상태 업데이트 헬퍼"""
    if task_id not in processing_tasks:
        processing_tasks[task_id] = {"id": task_id, "created_at": time.time()}
    
    processing_tasks[task_id].update({
        "status": status,
        "message": message,
        "updated_at": time.time(),
        **kwargs
    })
    print(f" ⏱ [Task {task_id}] {status}: {message}")

def load_queue_state():
    """파일에서 큐 상태를 로드하지 않습니다. (인메모리 전용 모드)"""
    pass

def save_queue_state():
    """파일에 큐 상태를 저장하지 않습니다. (인메모리 전용 모드)"""
    pass

def get_graph_store():
    """Neo4j 그래프 스토어 싱글톤"""
    global _graph_store
    try:
        if _graph_store is None:
            from backend.graph_store import Neo4jGraphStore
            _graph_store = Neo4jGraphStore()
            _graph_store.connect()
        return _graph_store
    except Exception as e:
        print(f" 🔴 [Graph Store] 싱글톤 초기화 실패: {e}")
        return None

#  Document pipeline
try:
    from backend.document_pipeline import process_document
    from dataclasses import dataclass

    @dataclass
    class Chunk:
        text: str
        metadata: dict
        index: int = 0

    LANGGRAPH_AVAILABLE = True
    print(" Document pipeline 사용 가능")
except ImportError as e:
    LANGGRAPH_AVAILABLE = False
    print(f" Document pipeline 사용 불가: {e}")

try:
    import langchain
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

@asynccontextmanager
async def lifespan(app: FastAPI):
    global _graph_store
    # Startup
    # (인메모리 전용 모드: 파일 복구 로직 제거)
    requeued_count = 0

    print(f"🚀 채팅 워커(Worker) 가동 중... (복구된 작업: {requeued_count}개)")
    worker_task = asyncio.create_task(chat_worker())

    yield

    # Shutdown
    print("\n 서버 종료 중...")
    worker_task.cancel()
    try:
        await worker_task
    except asyncio.CancelledError:
        print(" 채팅 워커 종료됨")

    try:
        vector_store.close_client()
    except Exception as ve:
        print(f" Weaviate 연결 종료 실패: {ve}")

    if _graph_store:
        try:
            _graph_store.close()
            print(" Neo4j 연결 종료됨")
        except Exception as ge:
            print(f" Neo4j 연결 종료 실패: {ge}")

# ═══════════════════════════════════════════════════════════════════════════
# PDF 변환 유틸리티
# ═══════════════════════════════════════════════════════════════════════════

def md_to_pdf_binary(md_text: str) -> bytes:
    """
    마크다운을 PDF 바이너리로 변환
    - <!-- PAGE:n --> 마커를 페이지 분할(<pdf:nextpage />)로 변환
    - 프론트엔드 UI와 유사한 계층적 들여쓰기 및 스타일 적용
    - Pretendard 한글 폰트 적용
    """
    try:
        import markdown
        from xhtml2pdf import pisa
        import re
        
        # 1. 페이지 마커 변환
        md_text = re.sub(r'<!-- PAGE:\d+ -->', '<pdf:nextpage />', md_text)
        
        # 2. 계층적 구조 분석 및 HTML 변환 (프론트엔드 App.tsx 로직 모방)
        lines = md_text.split('\n')
        processed_elements = []
        global_depth = 0
        indent_increment = 12 # pt
        
        paragraph_buffer = []
        table_buffer = []
        in_table = False
        
        def flush_paragraph():
            if paragraph_buffer:
                text = " ".join(paragraph_buffer).strip()
                if text:
                    padding = global_depth * indent_increment
                    html_text = markdown.markdown(text, extensions=['extra', 'nl2br', 'sane_lists'])
                    processed_elements.append(f'<div style="padding-left: {padding}pt; margin-bottom: 6pt; font-size: 10pt; line-height: 1.8; color: #333;">{html_text}</div>')
                paragraph_buffer.clear()

        def flush_table():
            if table_buffer:
                table_md = "\n".join(table_buffer)
                html_table = markdown.markdown(table_md, extensions=['tables'])
                processed_elements.append(f'<div style="margin: 15pt 0;">{html_table}</div>')
                table_buffer.clear()

        for line in lines:
            trimmed = line.strip()
            
            # 페이지 분할 태그 처리
            if trimmed == '<pdf:nextpage />':
                flush_paragraph()
                flush_table()
                processed_elements.append('<pdf:nextpage />')
                in_table = False
                continue
            
            # 테이블 감지
            if trimmed.startswith('|'):
                if not in_table:
                    flush_paragraph()
                    in_table = True
                table_buffer.append(line)
                continue
            elif in_table and trimmed:
                table_buffer.append(line)
                continue
            elif in_table and not trimmed:
                flush_table()
                in_table = False
                continue

            if not trimmed:
                continue

            # 조항 번호 패턴 매칭
            section_match = re.match(r'^(\d+(?:\.\d+)*)\.?\s+(.+)', trimmed)
            
            # 페이지 번호 무시
            if section_match and re.search(r'of\s+\d+', section_match.group(2), re.I):
                continue

            if section_match:
                flush_paragraph()
                section_num = section_match.group(1)
                section_text = section_match.group(2)
                depth = section_num.count('.')
                global_depth = depth
                
                padding = depth * indent_increment
                display_text = f"{section_num} {section_text}"
                
                if depth == 0:
                    style = f"font-weight: bold; font-size: 12pt; margin-top: 40pt; margin-bottom: 8pt; color: #1a1a1a; border-bottom: 0.5pt solid #e9ecef; padding-bottom: 8pt; padding-left: {padding}pt;"
                else:
                    style = f"font-size: 11pt; margin-top: 20pt; margin-bottom: 8pt; color: #2c3e50; padding-left: {padding}pt;"
                
                processed_elements.append(f'<div style="{style}">{display_text}</div>')
            else:
                paragraph_buffer.append(trimmed)
        
        flush_paragraph()
        flush_table()
        
        final_body_html = "".join(processed_elements)

        # 3. HTML 템플릿 구성
        font_path = "/Users/soyeon/Library/Fonts/Pretendard-Regular.ttf"
        bold_font_path = "/Users/soyeon/Library/Fonts/Pretendard-Bold.ttf"
        
        html_template = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @font-face {{
                    font-family: 'Pretendard';
                    src: url('{font_path}');
                }}
                @font-face {{
                    font-family: 'Pretendard';
                    src: url('{bold_font_path}');
                    font-weight: bold;
                }}
                @page {{
                    size: a4;
                    margin: 2cm;
                    @frame footer {{
                        -pdf-frame-content: footer_content;
                        bottom: 1cm;
                        margin-left: 2cm;
                        margin-right: 2cm;
                        height: 1cm;
                    }}
                }}
                body {{
                    font-family: 'Malgun Gothic', 'Pretendard', sans-serif;
                    color: #333;
                }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 15pt 0;
                }}
                th, td {{ 
                    border: 0.5pt solid #ccc; 
                    padding: 8pt; 
                    text-align: left; 
                    font-size: 9pt;
                }}
                th {{ 
                    background-color: #252526; 
                    color: #22D142; 
                    font-weight: bold; 
                }}
                tr:nth-child(even) {{ background-color: #fcfcfc; }}
            </style>
        </head>
        <body>
            {final_body_html}
            <div id="footer_content" style="text-align: center; color: #999; font-size: 9pt;">
                Page <pdf:pagenumber>
            </div>
        </body>
        </html>
        """
        
        out = BytesIO()
        pisa_status = pisa.CreatePDF(html_template, dest=out, encoding='utf-8')
        
        if pisa_status.err:
            raise Exception("PDF 생성 중 오류 발생")
            
        return out.getvalue()
    except Exception as e:
        print(f"⚠ PDF 변환 실패: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"PDF 변환 중 오류가 발생했습니다: {str(e)}")

def md_to_docx_binary(md_text: str, title: str) -> bytes:
    """마크다운을 Word(.docx) 바이너리로 변환 (간단한 구조)"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # 마크다운을 줄 단위로 분리하여 간단하게 구현 (정밀한 라이브러리 대신 직접 처리)
        lines = md_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        out = BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        print(f"⚠ Word 변환 실패: {e}")
        raise HTTPException(500, f"Word 변환 중 오류가 발생했습니다: {str(e)}")

app = FastAPI(title="RAG Chatbot API", version="9.2.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


async def process_save_document_task(
    doc_name: str,
    content: str,
    collection: str,
    model: str,
    task_id: str
):
    """문서 수정 저장 백그라운드 작업"""
    start_time = time.time()
    update_task_status(task_id, "processing", f"'{doc_name}' 문서 수정을 시작합니다.")
    
    try:
        # 1. 문서 재분석 (파이프라인 재사용)
        content_bytes = content.encode('utf-8')
        model_path = resolve_model_path(model)
        embed_model = SentenceTransformer(model_path)
        
        # 파이프라인 실행
        result = await asyncio.to_thread(
            process_document,
            file_path=f"{doc_name}.md",
            content=content_bytes,
            use_llm_metadata=True, # 메타데이터 및 버전 추출을 위해 활성화
            embed_model=embed_model
        )
        
        if not result.get("success"):
            raise Exception(f"🔴 분석 실패: {result.get('errors')}")
            
        final_version = result.get("version", "1.0")
        chunks_data = result["chunks"]
        doc_id = result.get("doc_id", doc_name)
        
        update_task_status(task_id, "processing", f"분석 완료 (버전 {final_version}). DB 동기화 중...", doc_id=doc_id)
        
        # 2. 기존 검색 데이터 삭제 (Overwrite 정제)
        await asyncio.to_thread(vector_store.delete_by_doc_name, doc_name, collection_name=collection)
        
        try:
            graph = get_graph_store()
            if graph and graph.test_connection():
                sop_id = doc_id
                if not re.search(r'[A-Z]+-[A-Z]+-\d+', sop_id):
                    sop_match = re.search(r'([A-Z]+-[A-Z]+-\d+)', doc_name, re.IGNORECASE)
                    if sop_match:
                        sop_id = sop_match.group(1).upper()
                
                await asyncio.to_thread(graph.delete_document, sop_id)
        except Exception as ge:
            print(f"  ⚠ Neo4j 삭제 실패 (무시): {ge}")

        # 3. RDB 신규 버전 저장
        doc_id_db = await asyncio.to_thread(
            sql_store.save_document,
            doc_name=doc_name,
            content=content,
            doc_type="text/markdown",
            version=final_version
        )
        
        if doc_id_db:
            batch_chunks = [
                {
                    "clause": c["metadata"].get("clause_id"),
                    "content": c["text"],
                    "metadata": c["metadata"]
                }
                for c in chunks_data
            ]
            await asyncio.to_thread(sql_store.save_chunks_batch, doc_id_db, batch_chunks)
        
        # 4. 벡터 DB 재업로드
        texts = [c["text"] for c in chunks_data]
        metadatas = [
            {
                **c["metadata"],
                "chunk_method": "article",
                "model": model,
                "pipeline_version": "edit-save-v2.0",
            }
            for c in chunks_data
        ]
        
        await asyncio.to_thread(
            vector_store.add_documents,
            texts=texts,
            metadatas=metadatas,
            collection_name=collection,
            model_name=model_path
        )
        
        # 5. 그래프 DB 재업로드
        try:
            graph = get_graph_store()
            if graph and graph.test_connection():
                await asyncio.to_thread(_upload_to_neo4j_from_pipeline, graph, result, f"{doc_name}.md")
        except Exception as ge:
            print(f"  ⚠ Neo4j 업로드 실패 (무시): {ge}")
            
        elapsed = round(time.time() - start_time, 2)
        update_task_status(task_id, "completed", f"문서 수정이 완료되었습니다. ({elapsed}초)", doc_name=doc_name, doc_id=doc_name, version=final_version)
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        update_task_status(task_id, "error", f"저장 중 오류 발생: {str(e)}", doc_name=doc_name)

@app.post("/rag/document/save")
async def save_document_content(request: SaveDocRequest, background_tasks: BackgroundTasks):
    """
    수정된 문서 내용을 저장하고 DB 동기화 (비동기)
    """
    task_id = f"save_{uuid.uuid4().hex[:8]}"
    update_task_status(task_id, "waiting", f"'{request.doc_name}' 수정 저장 요청이 접수되었습니다.", doc_name=request.doc_name)
    
    background_tasks.add_task(
        process_save_document_task,
        doc_name=request.doc_name,
        content=request.content,
        collection=request.collection,
        model=request.model,
        task_id=task_id
    )
    
    return {
        "success": True,
        "message": "문서 수정이 시작되었습니다.",
        "task_id": task_id,
        "doc_name": request.doc_name
    }

@app.post("/rag/upload-docx")
async def upload_docx(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    doc_name: str = Form(...),
    collection: str = Form("documents"),
    model: str = Form("multilingual-e5-small"),
    version: str = Form("1.0")
):
    """
    DOCX 파일 업로드 및 분석 (비동기)
    """
    try:
        content = await file.read()
        filename = file.filename
        
        task_id = f"upload_docx_{uuid.uuid4().hex[:8]}"
        # doc_name을 filename 대신 사용하여 알림에서 문서 ID가 보이게 함
        update_task_status(task_id, "waiting", f"'{doc_name}' 업로드 요청이 접수되었습니다.", doc_name=doc_name, filename=filename)

        # 기존 업로드 태스크 재사용 (DOCX도 동일 파이프라인 사용 가능하도록 유도)
        background_tasks.add_task(
            process_upload_task,
            filename=doc_name, # 파이프라인에서 doc_id로 사용됨
            content=content,
            collection=collection,
            chunk_size=500, # 기본값
            chunk_method="article",
            model=model,
            overlap=50,
            use_langgraph=True,
            use_llm_metadata=True,
            task_id=task_id,
            version=version,
        )

        return {
            "success": True,
            "message": f"'{doc_name}' 문서의 업로드가 시작되었습니다.",
            "task_id": task_id,
            "doc_name": doc_name
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"🔴 DOCX 업로드 요청 실패: {str(e)}")



# ═══════════════════════════════════════════════════════════════════════════
# 헬퍼 함수
# ═══════════════════════════════════════════════════════════════════════════

def resolve_model_path(model: str) -> str:
    """모델 프리셋 → 전체 경로"""
    return PRESET_MODELS.get(model, model)


def format_context(results: List[Dict]) -> str:
    """검색 결과 → 컨텍스트 문자열 (메타데이터 포함)"""
    context_parts = []
    
    for i, r in enumerate(results, 1):
        meta = r.get("metadata", {})
        text = r.get("text", "")
        similarity = r.get("similarity", 0)
        
        #  v9.2: 개선된 출처 표시
        doc_id = meta.get("doc_id", "")
        section_path = meta.get("section_path", "")
        page = meta.get("page", "")
        article_num = meta.get("article_num", "")
        
        # 출처 헤더 구성
        source_parts = []
        if doc_id:
            source_parts.append(f"[{doc_id}]")
        if section_path:
            source_parts.append(f"> {section_path}")
        if page:
            source_parts.append(f"(p.{page})")
        if similarity:
            source_parts.append(f"관련도: {similarity:.0%}")
        
        source_header = " ".join(source_parts) if source_parts else f"[문서 {i}]"
        
        context_parts.append(f"{source_header}\n{text}")
    
    return "\n\n---\n\n".join(context_parts)


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - 기본
# ═══════════════════════════════════════════════════════════════════════════

@app.get("/")
def root():
    return {
        "message": "RAG Chatbot API v9.2",
        "features": [
            "LangGraph 파이프라인",
            "페이지 번호 추적",
            "Parent-Child 계층",
            "Question 추적 (Neo4j)",
            "Weaviate + Neo4j 동기화 삭제"
        ],
        "endpoints": {
            "upload": "/rag/upload",
            "search": "/rag/search",
            "chat": "/chat",
            "ask": "/rag/ask",
            "graph": "/graph/*"
        },
        "langgraph_enabled": LANGGRAPH_AVAILABLE and USE_LANGGRAPH
    }


@app.get("/health")
def health():
    return {
        "status": "healthy",
        "cuda": torch.cuda.is_available(),
        "device": device,
        "ollama": OllamaLLM.is_available(),
        "langgraph": LANGGRAPH_AVAILABLE
    }


@app.get("/models/embedding")
def list_embedding_models():
    return {
        "presets": PRESET_MODELS,
        "specs": vector_store.EMBEDDING_MODEL_SPECS,
        "compatible": vector_store.filter_compatible_models()
    }


@app.get("/models/llm")
def list_llm_models():
    available_ollama = []
    if OllamaLLM.is_available():
        available_ollama = OllamaLLM.list_models()
    return {
        "ollama": {"presets": OLLAMA_MODELS, "available": available_ollama},
        "huggingface": HUGGINGFACE_MODELS
    }


# ═══════════════════════════════════════════════════════════════════════════
#  API 엔드포인트 - 업로드 (LangGraph v9.2)
# ═══════════════════════════════════════════════════════════════════════════

def process_upload_task(
    filename: str,
    content: bytes,
    collection: str,
    chunk_size: int,
    chunk_method: str,
    model: str,
    overlap: int,
    use_langgraph: bool,
    use_llm_metadata: bool,
    task_id: str,
    version: Optional[str] = None,
):
    """
    문서 업로드 처리 (배경 작업)
    - 기존 업로드 파이프라인을 그대로 사용
    """
    start_time = time.time()
    update_task_status(task_id, "processing", f"'{filename}' 분석을 시작합니다.")

    try:
        print(f"\n{'='*70}", flush=True)
        print(f"문서 업로드 처리 시작: {filename}", flush=True)
        print(f"{'='*70}\n", flush=True)

        # ========================================
        # 문서 파싱
        # ========================================
        print(f"[1단계] 문서 파싱", flush=True)
        print(f"  파이프라인: PDF 조항 v2.0", flush=True)
        print(f"  LLM 메타데이터: {'🟢 활성' if use_llm_metadata else '비활성'}", flush=True)
        if use_llm_metadata:
            print(f"  LLM 모델: gpt-4o", flush=True)
        print("", flush=True)

        model_path = resolve_model_path(model)
        embed_model = SentenceTransformer(model_path)

        result = process_document(
            file_path=filename,
            content=content,
            use_llm_metadata=use_llm_metadata,
            embed_model=embed_model
        )

        if not result.get("success"):
            errors = result.get("errors", ["알 수 없는 오류"])
            raise HTTPException(400, f"🔴 문서 처리 실패: {errors}")

        chunks_data = result["chunks"]
        if not chunks_data:
            raise HTTPException(400, "🔴 텍스트 추출 실패")

        from dataclasses import dataclass
        @dataclass
        class Chunk:
            text: str
            metadata: dict
            index: int = 0

        chunks = [Chunk(text=c["text"], metadata=c["metadata"], index=c["index"]) for c in chunks_data]
        doc_id = result.get("doc_id")
        doc_title = result.get("doc_title")
        pipeline_version = "pdf-clause-v2.0"

        print(f"  🟢 파싱 완료", flush=True)
        print(f"     • ID: {doc_id}", flush=True)
        print(f"     • 제목: {doc_title}", flush=True)
        print(f"     • 조항: {result.get('total_clauses')}개", flush=True)
        print(f"     • 청크: {len(chunks)}개\n", flush=True)
        
        update_task_status(task_id, "processing", f"파싱 완료 ({len(chunks)}개 청크). 벡터 DB 저장 중...", doc_id=doc_id)
        
        # ========================================
        # Weaviate 벡터 저장
        # ========================================
        print(f"[2단계] Weaviate 벡터 저장", flush=True)

        texts = [c.text for c in chunks]
        metadatas = [
            {
                **c.metadata,
                "chunk_method": chunk_method,
                "model": model,
                "pipeline_version": pipeline_version,
            }
            for c in chunks
        ]

        vector_store.add_documents(
            texts=texts,
            metadatas=metadatas,
            collection_name=collection,
            model_name=model_path
        )
        print(f"  🟢 저장 완료: {len(chunks)}개 청크\n", flush=True)
        update_task_status(task_id, "processing", "벡터 DB 저장 완료. PostgreSQL 및 그래프 DB 저장 중...", doc_id=doc_id)
        
        # ========================================
        # PostgreSQL 문서 저장
        # ========================================
        print(f"[3단계] PostgreSQL 저장", flush=True)

        try:
            # PDF에서 추출한 원본 텍스트 그대로 사용 (조항 번호 포함)
            original_text = result.get("markdown", "")
            if not original_text:
                # fallback: 청크들을 합침
                original_text = "\n\n".join([c.text for c in chunks])

            # 파이프라인에서 추출된 버전 또는 사용자 입력 버전 결정
            final_version = version or result.get("version", "1.0")
            
            if final_version != "1.0":
                print(f"     [추출] 최종 결정된 버전: {final_version}", flush=True)

            doc_id_db = sql_store.save_document(
                doc_name=doc_id,
                content=original_text,  # PDF 원본 텍스트 그대로 저장
                doc_type=filename.split('.')[-1] if '.' in filename else None,
                version=final_version
            )

            # 원본 PDF를 S3에 저장
            if filename.lower().endswith('.pdf'):
                try:
                    get_s3_store().upload_pdf(doc_id, final_version, content)
                    print(f"  🟢 원본 PDF S3 저장 완료: {doc_id}/v{final_version}", flush=True)
                except Exception as e:
                    print(f"  🟡 S3 PDF 저장 실패 (무시): {e}", flush=True)

            if doc_id_db and chunks:
                batch_chunks = [
                    {
                        "clause": c.metadata.get("clause_id"),
                        "content": c.text,
                        "metadata": c.metadata
                    }
                    for c in chunks
                ]
                sql_store.save_chunks_batch(doc_id_db, batch_chunks)
                print(f"  🟢 저장 완료: 문서 + {len(chunks)}개 청크\n", flush=True)
            else:
                print(f"  🔴 저장 실패: DB 저장에 실패했습니다 (ID 생성 불가)\n", flush=True)
        except Exception as sql_err:
            print(f"  🔴 저장 실패: {sql_err}\n", flush=True)

        # ========================================
        # Neo4j 그래프 저장
        # ========================================
        print(f"[4단계] Neo4j 그래프 저장", flush=True)
        graph_uploaded = False
        graph_sections = 0

        try:
            from backend.graph_store import Neo4jGraphStore

            graph = get_graph_store()
            if graph.test_connection():
                _upload_to_neo4j_from_pipeline(graph, result, filename)
                graph_uploaded = True
                stats = graph.get_graph_stats()
                graph_sections = stats.get("sections", 0)
                print(f"  🟢 저장 완료: {graph_sections}개 섹션\n", flush=True)
        except Exception as graph_error:
            # [디버그 로그 보강] 연결 실패 시 구체적인 에러 메시지 출력
            print(f"  🔴 Neo4j 연결 실패: {graph_error}", flush=True)
            import traceback
            traceback.print_exc()
            print(f"  ⚠ 그래프 연동을 건너뛰고 계속 진행합니다.\n", flush=True)
        
        # 완료
        # ========================================
        elapsed = round(time.time() - start_time, 2)
        update_task_status(task_id, "completed", f"문서 업로드가 완료되었습니다. ({elapsed}초)", doc_id=doc_id, doc_name=doc_id, filename=filename, version=final_version)

        print(f"{'='*70}", flush=True)
        print(f"🟢 업로드 처리 완료 ({elapsed}초)", flush=True)
        print(f"{'='*70}\n", flush=True)

    except HTTPException as e:
        print(f"🔴 업로드 처리 실패: {e.detail}", flush=True)
        update_task_status(task_id, "error", f"업로드 처리 실패: {e.detail}", filename=filename)
        return
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"🔴 업로드 처리 실패: {str(e)}", flush=True)
        update_task_status(task_id, "error", f"알 수 없는 오류 발생: {str(e)}", filename=filename)
        return


@app.post("/rag/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    collection: str = Form("documents"),
    chunk_size: int = Form(DEFAULT_CHUNK_SIZE),
    chunk_method: str = Form(DEFAULT_CHUNK_METHOD),
    model: str = Form("multilingual-e5-small"),
    overlap: int = Form(DEFAULT_OVERLAP),
    use_langgraph: bool = Form(True),
    use_llm_metadata: bool = Form(True),
    version: Optional[str] = Form(None),
):
    """
    문서 업로드 요청을 즉시 접수하고, 실제 처리는 배경 작업으로 수행.
    """
    try:
        content = await file.read()
        filename = file.filename

        print(f"\n{'='*70}", flush=True)
        print(f"문서 업로드 요청 접수: {filename}", flush=True)
        print("  처리 방식: 비동기 (Background Tasks)", flush=True)
        print(f"{'='*70}\n", flush=True)

        task_id = f"upload_{uuid.uuid4().hex[:8]}"
        update_task_status(task_id, "waiting", f"'{filename}' 업로드 요청이 접수되었습니다.", filename=filename)

        background_tasks.add_task(
            process_upload_task,
            filename=filename,
            content=content,
            collection=collection,
            chunk_size=chunk_size,
            chunk_method=chunk_method,
            model=model,
            overlap=overlap,
            use_langgraph=use_langgraph,
            use_llm_metadata=use_llm_metadata,
            task_id=task_id,
            version=version,
        )

        return {
            "success": True,
            "message": f"'{filename}' 문서의 업로드가 시작되었습니다.",
            "task_id": task_id,
            "filename": filename,
            "processing_mode": "background",
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"🔴 요청 실패: {str(e)}")


def _upload_to_neo4j_from_pipeline(graph, result: dict, filename: str):
    """새 파이프라인 결과를 Neo4j에 업로드 (간소화)"""
    from backend.graph_store import upload_document_to_graph
    upload_document_to_graph(graph, result, filename)


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - 검색
# ═══════════════════════════════════════════════════════════════════════════

# /rag/search 엔드포인트 제거됨 (Agent가 내부 수행)


# /rag/search/advanced 엔드포인트 제거됨


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - 챗봇
# ═══════════════════════════════════════════════════════════════════════════

async def process_chat_request(request: ChatRequest) -> Dict:
    """실제 에이전트 질의 처리 본문 (워커/동기 엔드포인트 공용)."""
    print(f" [Agent] 요청 수신: {request.message}")

    # Agent 초기화
    init_agent_tools(vector_store, get_graph_store(), sql_store)

    # 세션 ID 선할당 (None 방지)
    session_id = request.session_id or str(uuid.uuid4())

    # 롱텀 메모리: 이전 대화 기록 + 유사 기억 로드
    chat_history = []
    user_id = request.user_id
    query_embedding = None

    if user_id:
        try:
            if request.session_id:
                chat_history = sql_store.get_conversation_history_by_session(
                    user_id,
                    request.session_id,
                    limit=100
                )
                print(f"  🧠 [Memory] 사용자 {user_id}, 세션 {request.session_id} 기록 {len(chat_history)}건 로드")
            else:
                chat_history = sql_store.get_conversation_history(user_id, limit=10)
                print(f"  🧠 [Memory] 사용자 {user_id} 최신 기록 {len(chat_history)}건 로드 (세션 미지정)")

            print(f"  🧠 [Memory] 사용자 {user_id} 컨텍스트 로드 완료")
        except Exception as e:
            print(f"  ⚠️ [Memory] 조회 실패: {e}")
            import traceback
            traceback.print_exc()

    effective_model = resolve_effective_llm_model(request.llm_model)
    print(f" [LLM] effective_model={effective_model} (requested={request.llm_model})")

    response = await asyncio.to_thread(
        run_agent,
        query=request.message,
        session_id=session_id,
        model_name=effective_model,
        chat_history=chat_history
    )

    answer = response.get("answer") or ""

    # 롱텀 메모리: 새로운 대화 저장 (세션 기준)
    if user_id and answer:
        try:
            target_session_id = session_id

            query_embedding = None
            try:
                # 무거운 모델 로드는 스레드에서 실행
                def _get_embedding():
                    embed_model = SentenceTransformer("intfloat/multilingual-e5-small")
                    return embed_model.encode(request.message).tolist()
                
                query_embedding = await asyncio.to_thread(_get_embedding)
            except Exception as embed_error:
                print(f"  ⚠️ [Memory] 임베딩 생성 실패: {embed_error}")

            await asyncio.to_thread(
                sql_store.save_memory,
                request.message,
                answer,
                user_id,
                target_session_id,
                embedding=query_embedding
            )
            print(f"  💾 [Memory] 세션 {target_session_id}에 대화 내용 저장 완료")
        except Exception as e:
            print(f"  ⚠️ [Memory] 저장 실패: {e}")
            import traceback
            traceback.print_exc()

    # LLM as a Judge 평가
    evaluation_scores = None
    error_patterns = ["오류가 발생", "에러", "실패", "Error", "Exception", "찾을 수 없", "준비하지 못", "로딩 에러"]
    is_error_message = any(pattern in answer for pattern in error_patterns)

    try:
        from backend.evaluation import AgentEvaluator
        if len(answer) >= 20 and not is_error_message:
            # 평가기(AgentEvaluator)는 내부적으로 동기 통신을 수행하므로 별도 스레드에서 실행
            def _evaluate():
                evaluator = AgentEvaluator(judge_model="gpt-4o", sql_store=sql_store)
                context = response.get("agent_log", {}).get("context", "")
                if isinstance(context, list):
                    context = "\n\n".join(context)

                return evaluator.evaluate_single(
                    question=request.message,
                    answer=answer,
                    context=context,
                    metrics=["faithfulness", "groundness", "relevancy", "correctness"]
                )
            evaluation_scores = await asyncio.to_thread(_evaluate)
    except ImportError:
        print("평가 모듈 사용 불가 (선택적 기능)")
    except Exception as eval_error:
        print(f"평가 실행 실패 (계속 진행): {eval_error}")
        evaluation_scores = None

    return {
        "session_id": session_id,
        "answer": answer,
        "sources": [],
        "agent_log": response,
        "evaluation_scores": evaluation_scores
    }


async def chat_worker():
    """큐에서 질문을 하나씩 꺼내어 순차적으로 답변을 생성하는 워커."""
    while True:
        request_id, ticket, kind, payload = await chat_queue.get()
        try:
            print(f" 🚀 [Chat Worker] 처리 시작: {request_id} (ticket: {ticket}, kind: {kind})")
            
            load_queue_state()
            if request_id in chat_results:
                chat_results[request_id]["status"] = "processing"
            save_queue_state()

            if kind == "rag":
                req = ChatRequest(**payload)
                result_data = await process_chat_request(req)

            elif kind == "agent":
                req = AgentRequest(**payload)
                # 에이전트 도구 초기화 (동작 확인용으로 남겨둠, thread 내부에서 수행 권장)
                from backend.agent import init_agent_tools, run_agent
                
                # 비차단/멀티스레드 호출로 변경 (동기 함수인 run_agent가 루프를 막지 않게 함)
                result = await asyncio.to_thread(
                    run_agent,
                    query=req.message,
                    session_id=req.session_id,
                    model_name=req.llm_model,
                    embedding_model=resolve_model_path(req.embedding_model)
                )

                answer = result.get("answer", "")
                reasoning = result.get("reasoning", "")
                
                # 본문이 비어있으면 사고 과정(Think)을 답변으로 사용 (전처리 로직 동일화)
                if not answer and reasoning:
                    answer = f"[AI 분석 리포트]\n\n{reasoning}"

                result_data = {
                    "session_id": req.session_id,
                    "answer": answer,
                    "reasoning": reasoning,
                    "tool_calls": result.get("tool_calls", []),
                    "success": result.get("success", False),
                    "mode": "langgraph" if (req.use_langgraph and LANGGRAPH_AGENT_AVAILABLE) else "simple"
                }
            else:
                raise ValueError(f"Unknown kind: {kind}")
            
            load_queue_state()
            if request_id in chat_results:
                chat_results[request_id].update({
                    "status": "completed",
                    "result": result_data
                })
            save_queue_state()
            
            print(f" ✅ [Chat Worker] 처리 완료: {request_id}")
        except Exception as e:
            print(f" 🔴 [Chat Worker] 에러: {e}")
            import traceback
            traceback.print_exc()
            load_queue_state()
            if request_id in chat_results:
                chat_results[request_id].update({
                    "status": "error",
                    "error": str(e)
                })
            save_queue_state()
        finally:
            load_queue_state()
            # pending에서 제거(request_id 기준)
            # pending에서 제거(request_id 기준) - 인메모리 리스트 객체 보존을 위해 슬라이스 할당 사용
            chat_pending[:] = [x for x in chat_pending if x["request_id"] != request_id]
            save_queue_state()
            chat_queue.task_done()

def _extract_user_id_from_auth_header(auth_header: Optional[str]) -> Optional[int]:
    """Authorization 헤더의 JWT에서 user_id를 추출."""
    if not auth_header or not auth_header.startswith("Bearer "):
        return None
    token = auth_header.split(" ", 1)[1].strip()
    if not token:
        return None
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("user_id")
        return int(user_id) if user_id is not None else None
    except Exception:
        return None


async def enqueue_job(kind: str, payload: dict) -> dict:
    """
    kind: "rag" | "agent"
    payload: worker가 처리할 요청 데이터(직렬화 가능한 dict)
    """
    global next_ticket, chat_pending

    request_id = str(uuid.uuid4())

    async with ticket_lock:
        next_ticket += 1
        ticket = next_ticket

    load_queue_state()

    # 대기열 엔트리 추가
    chat_pending.append({"request_id": request_id, "ticket": ticket, "kind": kind})

    # 상태 저장(클라이언트 status 조회용)
    chat_results[request_id] = {
        "status": "waiting",
        "ticket": ticket,
        "kind": kind,
        "result": None,
        "request": payload,
    }
    save_queue_state()

    await chat_queue.put((request_id, ticket, kind, payload))

    # position(현재 대기 순서) = pending에서 ticket 작은 애들 수 + 1
    position = sum(1 for x in chat_pending if x["ticket"] < ticket) + 1

    return {
        "success": True,
        "request_id": request_id,
        "ticket": ticket,
        "status": "waiting",
        "position": position,
        "message": f"질문이 대기열에 등록되었습니다. (현재 대기 순번: {position}번째)",
    }


@app.post("/chat")
async def chat(chat_request: ChatRequest, http_request: Request):
    """순차 대기열 채팅 엔드포인트."""
    session_id = chat_request.session_id or str(uuid.uuid4())
    user_id = chat_request.user_id
    if user_id is None:
        auth_header = http_request.headers.get("Authorization")
        user_id = _extract_user_id_from_auth_header(auth_header)

    effective_model = resolve_effective_llm_model(chat_request.llm_model)
    queued_request = chat_request.model_copy(update={
        "session_id": session_id,
        "user_id": user_id,
        "llm_model": effective_model
    })

    # 핵심: 큐 등록은 공통 함수 사용
    return await enqueue_job(
        kind="rag",
        payload=queued_request.model_dump()
    )


@app.get("/chat/status/{request_id}")
async def get_chat_status(request_id: str):
    """채팅 작업의 상태와 대기 순번 조회."""
    load_queue_state()
    
    if request_id not in chat_results:
        raise HTTPException(404, "요청 ID를 찾을 수 없습니다.")

    status_data = chat_results[request_id].copy()

    # ticket 상시 포함 보장 (이미 포함되어 있음)
    # status가 'waiting'일 때만 position 계산하여 노출
    if status_data.get("status") == "waiting":
        mine = next((x for x in chat_pending if x["request_id"] == request_id), None)
        if mine:
            my_ticket = mine["ticket"]
            # position(현재 대기 순서) = pending에서 내 티켓보다 작은 애들 수 + 1
            status_data["position"] = sum(1 for x in chat_pending if x["ticket"] < my_ticket) + 1
        else:
            status_data.pop("position", None)
    else:
        # processing, completed 등에서는 position 제외
        status_data.pop("position", None)

    return status_data


@app.get("/chat/history/{session_id}")
def get_chat_history(session_id: str):
    """대화 히스토리 조회"""
    history = chat_histories.get(session_id, [])
    return {"session_id": session_id, "history": history, "count": len(history)}


@app.delete("/chat/history/{session_id}")
def clear_chat_history(session_id: str):
    """대화 히스토리 삭제"""
    if session_id in chat_histories:
        del chat_histories[session_id]
        return {"success": True, "message": f"세션 {session_id} 삭제됨"}
    return {"success": False, "message": "세션을 찾을 수 없음"}


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - LLM
# ═══════════════════════════════════════════════════════════════════════════

@app.post("/llm/generate")
def generate_llm(request: LLMRequest):
    """LLM 직접 호출"""
    try:
        response = get_llm_response(
            prompt=request.prompt,
            llm_model=request.model,
            llm_backend=request.backend,
            max_tokens=request.max_tokens,
            temperature=request.temperature
        )
        return {"response": response, "model": request.model, "backend": request.backend}
    except Exception as e:
        raise HTTPException(500, f"LLM 호출 실패: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - 문서 관리
# ═══════════════════════════════════════════════════════════════════════════

def extract_document_category(doc_id: str) -> str:
    """문서 ID에서 카테고리 추출 (SOP, WI, FRM)"""
    if not doc_id:
        return "기타"
    doc_id_upper = doc_id.upper()
    if "SOP" in doc_id_upper:
        return "SOP"
    elif "WI" in doc_id_upper:
        return "WI"
    elif "FRM" in doc_id_upper or "FORM" in doc_id_upper:
        return "FRM"
    else:
        return "기타"


@app.get("/rag/documents")
async def list_documents(collection: str = "documents"):
    """문서 목록 (RDB + S3 DOCX 병합 조회)"""
    try:
        # SQL Store에서 모든 문서 조회 (비동기 처리)
        all_docs = await asyncio.to_thread(sql_store.get_all_documents)

        def _normalize_doc_type(v: Optional[str]) -> str:
            t = (v or "").lower()
            if "pdf" in t:
                return "pdf"
            if "docx" in t:
                return "docx"
            return t or "other"

        # 문서별로 그룹화 (같은 문서의 여러 버전)
        grouped: Dict[str, List[Dict]] = {}
        for doc in all_docs:
            doc_name = doc.get("doc_name")
            if not doc_name:
                continue
            grouped.setdefault(doc_name, []).append(doc)

        docs_out: List[Dict] = []
        for doc_name, versions in grouped.items():
            # 동일 문서명에서 pdf가 하나라도 있으면 pdf 계열만 대표 후보로 사용
            pdf_versions = [d for d in versions if _normalize_doc_type(d.get("doc_type")) == "pdf"]
            candidates = pdf_versions if pdf_versions else versions

            def _sort_key(d: Dict):
                created = d.get("created_at")
                if hasattr(created, "timestamp"):
                    created_key = created.timestamp()
                else:
                    created_key = 0
                return (created_key, str(d.get("version") or "0"))

            # 대표 문서는 created_at 최신 우선 (없으면 version 문자열로 보조)
            selected = max(candidates, key=_sort_key)

            docs_out.append({
                "doc_id": doc_name,
                "doc_name": doc_name,
                "doc_type": selected.get("doc_type"),
                "doc_format": _normalize_doc_type(selected.get("doc_type")),
                "doc_category": extract_document_category(doc_name),
                "version": selected.get("version"),
                "created_at": selected.get("created_at"),
                "latest_version": selected.get("version"),
                "source": "rdb",
            })

        # S3 DOCX 문서 병합: 동일 doc_name이 있어도 숨기지 않고 표시
        try:
            s3_docs = await asyncio.to_thread(get_s3_store().list_docx_documents)
            for s3_doc in s3_docs:
                doc_name = s3_doc.get("doc_name")
                docs_out.append({
                    "doc_id": doc_name,
                    "doc_name": doc_name,
                    "doc_type": "docx",
                    "doc_format": "docx",
                    "doc_category": extract_document_category(doc_name),
                    "version": s3_doc.get("version"),
                    "created_at": s3_doc.get("created_at"),
                    "latest_version": s3_doc.get("latest_version"),
                    "source": "s3",
                })
        except Exception as s3_err:
            print(f"S3 DOCX 목록 병합 실패(무시): {s3_err}")

        return {"documents": docs_out, "collection": collection}
    except Exception as e:
        print(f"문서 목록 조회 실패: {e}")
        return {"documents": [], "collection": collection}


@app.get("/rag/doc-names")
async def list_doc_names():
    """모든 문서 이름 목록 조회 (RDB doc_name 테이블)"""
    try:
        doc_names = await asyncio.to_thread(sql_store.list_doc_names)
        return {"doc_names": doc_names}
    except Exception as e:
        print(f"문서 이름 목록 조회 실패: {e}")
        return {"doc_names": []}


@app.get("/rag/document/{doc_name}/pdf-url")
def get_pdf_presigned_url(doc_name: str, version: Optional[str] = None):
    """PDF 열람 URL 반환 - S3 presigned URL 우선, 없으면 download 엔드포인트 사용"""
    try:
        # 버전이 없으면 최신 버전 사용
        if not version:
            versions = sql_store.get_document_versions(doc_name)
            if not versions:
                raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name}")
            version = versions[0]['version']

        # S3에서 원본 PDF 시도
        try:
            store = get_s3_store()
            if store.pdf_exists(doc_name, version):
                url = store.get_pdf_presigned_url(doc_name, version)
                return {"url": url, "source": "s3", "doc_name": doc_name, "version": version}
        except Exception as s3_err:
            print(f"  S3 PDF 조회 실패 (download로 폴백): {s3_err}")

        # 폴백: 백엔드 download 엔드포인트 사용 (프론트에서 auth 헤더 필요)
        return {"url": None, "source": "download", "doc_name": doc_name, "version": version}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"PDF URL 조회 실패: {str(e)}")


@app.get("/rag/document/{doc_name}/versions")
async def get_document_versions(doc_name: str):
    """문서 버전 목록 조회"""
    try:
        versions = await asyncio.to_thread(sql_store.get_document_versions, doc_name)
        if not versions:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name}")
        return {
            "doc_name": doc_name,
            "versions": versions,
            "count": len(versions)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"버전 조회 실패: {str(e)}")


@app.get("/rag/document/{doc_name}/compare")
async def compare_versions(doc_name: str, v1: str, v2: str):
    """두 버전 간의 조항 단위 차이 비교"""
    try:
        diffs = await asyncio.to_thread(sql_store.get_clause_diff, doc_name, v1, v2)

        # 에러가 있는지 확인
        if diffs and isinstance(diffs[0], dict) and 'error' in diffs[0]:
            raise HTTPException(400, diffs[0]['error'])

        return {
            "doc_name": doc_name,
            "v1": v1,
            "v2": v2,
            "diffs": diffs,
            "total_changes": len(diffs)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"버전 비교 실패: {str(e)}")


@app.get("/processing/status/{task_id}")
async def get_processing_status(task_id: str):
    """백그라운드 작업 상태 조회"""
    if task_id not in processing_tasks:
        raise HTTPException(404, "작업 ID를 찾을 수 없습니다.")
    return processing_tasks[task_id]

@app.get("/processing/list")
async def list_processing_tasks():
    """현재 관리 중인 모든 백그라운드 작업 목록"""
    return list(processing_tasks.values())


@app.get("/rag/changes")
async def get_changes(limit: int = 50):
    """최근 문서 변경 이력 조회"""
    try:
        # SQLStore에서 모든 문서를 가져와서 최근 수정순으로 반환 (비동기 처리)
        all_docs = await asyncio.to_thread(sql_store.get_all_documents)
        changes = []
        for doc in all_docs[:limit]:
            changes.append({
                "id": str(doc.get('id')),
                "doc_id": doc.get('doc_name'),
                "change_type": "UPDATE" if doc.get('version') != "1.0" else "CREATE",
                "changed_at": doc.get('created_at').isoformat() if doc.get('created_at') else None,
                "changed_by": "System",
                "description": f"Version {doc.get('version')} saved."
            })
        return {"changes": changes, "count": len(changes)}
    except Exception as e:
        raise HTTPException(500, f"변경 이력 조회 실패: {str(e)}")


@app.get("/rag/document/{doc_name}/content")
async def get_document_content(doc_name: str, version: Optional[str] = None):
    """문서 전체 내용 조회"""
    try:
        doc = await asyncio.to_thread(sql_store.get_document_by_name, doc_name, version)
        if not doc:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name} (v{version or '최신'})")

        # 청크 조회
        chunks = await asyncio.to_thread(sql_store.get_chunks_by_document, doc['id'])

        return {
            "doc_name": doc_name,
            "version": doc.get('version', '1.0'),
            "doc_type": doc.get('doc_type'),
            "created_at": doc.get('created_at'),
            "content": doc.get('content', ''),
            "chunks": chunks,
            "chunk_count": len(chunks)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"문서 조회 실패: {str(e)}")


@app.get("/rag/document/{doc_name}/metadata")
async def get_document_metadata(doc_name: str, version: Optional[str] = None):
    """문서 메타데이터 조회"""
    try:
        doc = await asyncio.to_thread(sql_store.get_document_by_name, doc_name, version)
        if not doc:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name}")

        # 청크 통계
        chunks = await asyncio.to_thread(sql_store.get_chunks_by_document, doc['id'])

        # 조항 통계 (메타데이터에서 추출)
        clause_ids = set()
        for chunk in chunks:
            metadata = chunk.get('metadata', {})
            if isinstance(metadata, dict):
                clause_id = metadata.get('clause_id')
                if clause_id:
                    clause_ids.add(clause_id)

        return {
            "doc_name": doc_name,
            "version": doc.get('version', '1.0'),
            "doc_type": doc.get('doc_type'),
            "created_at": doc.get('created_at'),
            "chunk_count": len(chunks),
            "clause_count": len(clause_ids),
            "total_length": len(doc.get('content', '')),
            "clauses": sorted(list(clause_ids), key=lambda x: [int(n) if n.isdigit() else n for n in x.split('.')])
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"메타데이터 조회 실패: {str(e)}")


@app.get("/rag/document/download/{doc_name}")
async def download_document(
    doc_name: str, 
    version: Optional[str] = None,
    format: Literal["pdf", "docx", "md"] = "pdf"
):
    """문서를 다양한 파일 형식으로 다운로드"""
    try:
        # DB에서 문서 조회 (버전 명시 없으면 최신본)
        doc = await asyncio.to_thread(sql_store.get_document_by_name, doc_name, version)
        if not doc:
            raise HTTPException(404, "문서를 찾을 수 없습니다.")
            
        content = doc["content"] # text/markdown 원본
        doc_type = doc["doc_type"]
        ver = doc["version"]
        
        # 파일명 구성 (공백 제거)
        safe_filename = doc_name.replace(" ", "_").replace("/", "_")
        
        # 1. 마크다운(.md) 형식 요청
        if format == "md":
            filename = f"{safe_filename}_v{ver}.md"
            return Response(
                content=content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
            
        # 2. 워드(.docx) 형식 요청
        if format == "docx":
            filename = f"{safe_filename}_v{ver}.docx"
            docx_bytes = await asyncio.to_thread(md_to_docx_binary, content, doc_name)
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
            
        # 3. PDF 형식 요청 (기존 로직 유지)
        filename = f"{safe_filename}_v{ver}.pdf"
        
        # 3-1. 이미 PDF인 경우
        if doc_type == "pdf" and isinstance(content, bytes):
            return Response(
                content=content,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
            
        # 3-2. 텍스트(마크다운)인 경우 PDF로 변환
        pdf_bytes = await asyncio.to_thread(md_to_pdf_binary, content)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"다운로드 처리 중 오류 발생: {str(e)}")


@app.delete("/rag/document")
def delete_document(request: DeleteDocRequest):
    """
    문서 삭제 (RDB + Weaviate + Neo4j 전체 삭제 + 삭제 검증)
    """
    print(f"\n[DELETE] 단건 삭제 요청: doc={request.doc_name}, collection={request.collection}, neo4j={request.delete_from_neo4j}")
    result = _delete_document_everywhere(
        doc_name=request.doc_name,
        collection=request.collection,
        delete_from_neo4j=request.delete_from_neo4j,
    )
    print(f"[DELETE] 단건 삭제 결과: doc={request.doc_name}, success={result.get('success')}, details={result.get('details')}")
    return {
        "success": result.get("success", False),
        "doc_name": request.doc_name,
        "details": result.get("details", {}),
    }


@app.post("/rag/documents/delete-batch")
def delete_documents_batch(request: DeleteDocsBatchRequest):
    """
    문서 다건 삭제 (RDB + Weaviate + Neo4j 전체 삭제 + 삭제 검증)
    - 하나라도 삭제 검증에 실패하면 overall_success=False
    """
    doc_names = [d.strip() for d in request.doc_names if d and d.strip()]
    if not doc_names:
        raise HTTPException(400, "삭제할 문서가 없습니다.")

    # 중복 제거(순서 유지)
    unique_doc_names = list(dict.fromkeys(doc_names))
    print(f"\n[DELETE] 배치 삭제 요청: count={len(unique_doc_names)}, docs={unique_doc_names}, collection={request.collection}, neo4j={request.delete_from_neo4j}")
    results = []
    success_count = 0

    for doc_name in unique_doc_names:
        item = _delete_document_everywhere(
            doc_name=doc_name,
            collection=request.collection,
            delete_from_neo4j=request.delete_from_neo4j,
        )
        results.append({
            "doc_name": doc_name,
            **item,
        })
        if item.get("success"):
            success_count += 1
        else:
            print(f"[DELETE] 배치 항목 실패: doc={doc_name}, details={item.get('details')}")

    print(f"[DELETE] 배치 삭제 결과: success={success_count}/{len(unique_doc_names)}")

    return {
        "success": success_count == len(unique_doc_names),
        "requested_count": len(unique_doc_names),
        "deleted_count": success_count,
        "failed_count": len(unique_doc_names) - success_count,
        "results": results,
    }


def _delete_document_everywhere(doc_name: str, collection: str, delete_from_neo4j: bool = True) -> Dict:
    """문서 1건을 세 저장소에서 삭제하고 실제 삭제 여부를 검증한다."""
    result = {"rdb": None, "weaviate": None, "neo4j": None, "s3_docx": None}
    print(f"[DELETE] 처리 시작: doc={doc_name}")
    doc_id_match = re.search(r'(EQ-(?:SOP|WI|FRM)-\d+)', doc_name, re.IGNORECASE)
    normalized_doc_id = doc_id_match.group(1).upper() if doc_id_match else None
    candidate_names = [doc_name]
    if normalized_doc_id and normalized_doc_id not in candidate_names:
        candidate_names.insert(0, normalized_doc_id)
    print(f"[DELETE] 후보 키: {candidate_names}")

    # 1) RDB 삭제 + 검증
    try:
        for name in candidate_names:
            sql_store.delete_document_by_name(name)
        still_exists = any(sql_store.get_document_by_name(name) is not None for name in candidate_names)
        result["rdb"] = {"success": not still_exists, "exists_after_delete": still_exists, "candidates": candidate_names}
        print(f"[DELETE][RDB] doc={doc_name}, exists_after_delete={still_exists}, candidates={candidate_names}")
    except Exception as e:
        result["rdb"] = {"success": False, "error": str(e)}
        print(f"[DELETE][RDB] doc={doc_name}, error={e}")

    # 2) Weaviate 삭제 + 검증
    try:
        deleted_total = 0
        for name in candidate_names:
            delete_res = vector_store.delete_by_doc_name(doc_name=name, collection_name=collection)
            deleted_total += int(delete_res.get("deleted", 0))
        docs_after = vector_store.list_documents(collection_name=collection)
        remains = any(
            (d.get("doc_id") in candidate_names) or (d.get("doc_name") in candidate_names)
            for d in docs_after
        )
        result["weaviate"] = {
            "success": not remains,
            "deleted": deleted_total,
            "exists_after_delete": remains,
            "candidates": candidate_names,
        }
        print(f"[DELETE][Weaviate] doc={doc_name}, deleted={deleted_total}, exists_after_delete={remains}, candidates={candidate_names}")
    except Exception as e:
        result["weaviate"] = {"success": False, "error": str(e)}
        print(f"[DELETE][Weaviate] doc={doc_name}, error={e}")

    # 3) Neo4j 삭제 + 검증
    try:
        if not delete_from_neo4j:
            result["neo4j"] = {"success": True, "skipped": True}
            print(f"[DELETE][Neo4j] doc={doc_name}, skipped=True")
        else:
            graph = get_graph_store()
            if not graph.test_connection():
                result["neo4j"] = {"success": False, "error": "Neo4j 연결 실패"}
                print(f"[DELETE][Neo4j] doc={doc_name}, error=Neo4j 연결 실패")
            else:
                candidate_ids = candidate_names[:]

                for candidate in candidate_ids:
                    graph.delete_document(candidate)

                remains = any(graph.get_document(candidate) for candidate in candidate_ids)
                result["neo4j"] = {
                    "success": not remains,
                    "doc_ids": candidate_ids,
                    "exists_after_delete": remains,
                }
                print(f"[DELETE][Neo4j] doc={doc_name}, doc_ids={candidate_ids}, exists_after_delete={remains}")
    except Exception as e:
        result["neo4j"] = {"success": False, "error": str(e)}
        print(f"[DELETE][Neo4j] doc={doc_name}, error={e}")

    # 4) S3 DOCX 삭제 + 검증
    try:
        s3 = get_s3_store()
        deleted_count = 0
        for name in candidate_names:
            deleted_count += s3.delete_docx_versions(name)
        remains = any(s3.has_docx(name) for name in candidate_names)
        result["s3_docx"] = {
            "success": not remains,
            "deleted": deleted_count,
            "exists_after_delete": remains,
            "candidates": candidate_names,
        }
        print(f"[DELETE][S3-DOCX] doc={doc_name}, deleted={deleted_count}, exists_after_delete={remains}, candidates={candidate_names}")
    except Exception as e:
        result["s3_docx"] = {"success": False, "error": str(e)}
        print(f"[DELETE][S3-DOCX] doc={doc_name}, error={e}")

    success = (
        result["rdb"].get("success", False)
        and result["weaviate"].get("success", False)
        and result["neo4j"].get("success", False)
        and result["s3_docx"].get("success", False)
    )
    print(f"[DELETE] 처리 종료: doc={doc_name}, success={success}")
    return {"success": success, "details": result}


@app.get("/rag/collections")
def list_collections():
    """컬렉션 목록"""
    collections = vector_store.list_collections()
    return {"collections": [vector_store.get_collection_info(name) for name in collections]}


@app.delete("/rag/collection/{collection_name}")
def delete_collection(collection_name: str):
    """컬렉션 삭제"""
    return vector_store.delete_all(collection_name)


@app.get("/rag/supported-formats")
def get_supported_formats():
    """지원 포맷"""
    return {"supported_extensions": get_supported_extensions()}


@app.get("/rag/chunk-methods")
def get_chunk_methods():
    """청킹 방법"""
    return {"methods": get_available_methods()}


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - Neo4j 그래프
# ═══════════════════════════════════════════════════════════════════════════

@app.get("/graph/status")
def graph_status():
    """Neo4j 연결 상태"""
    try:
        graph = get_graph_store()
        connected = graph.test_connection()
        stats = graph.get_graph_stats() if connected else {}
        return {"connected": connected, "stats": stats}
    except Exception as e:
        return {"connected": False, "error": str(e)}


@app.post("/graph/init")
def graph_init():
    """Neo4j 스키마 초기화"""
    try:
        graph = get_graph_store()
        graph.init_schema()
        return {"success": True, "message": "스키마 초기화 완료"}
    except Exception as e:
        raise HTTPException(500, f"스키마 초기화 실패: {str(e)}")


@app.delete("/graph/clear")
def graph_clear():
    """Neo4j 모든 데이터 삭제"""
    try:
        graph = get_graph_store()
        graph.clear_all()
        return {"success": True, "message": "모든 데이터 삭제 완료"}
    except Exception as e:
        raise HTTPException(500, f"데이터 삭제 실패: {str(e)}")


@app.post("/graph/upload")
async def graph_upload_document(
    file: UploadFile = File(...),
    use_langgraph: bool = Form(True)
):
    """문서를 Neo4j 그래프로만 업로드"""
    try:
        content = await file.read()
        filename = file.filename
        
        if not LANGGRAPH_AVAILABLE:
            raise HTTPException(500, "LangGraph 모듈이 필요합니다.")
            
        result = process_document(filename, content, debug=True)
        if not result.get("success"):
            raise HTTPException(400, f"처리 실패: {result.get('errors')}")
        
        graph = get_graph_store()
        _upload_to_neo4j_from_pipeline(graph, result, filename)
        
        return {
            "success": True,
            "filename": filename,
            "doc_id": result.get("metadata", {}).get("doc_id"),
            "sections": len(result.get("sections", [])),
            "pipeline": "langgraph"
        }
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"그래프 업로드 실패: {str(e)}")


@app.get("/rag/documents")
async def list_documents(
    user_id: Optional[str] = None,
    doc_type: Optional[str] = None
):
    """업로드된 문서 목록 조회"""
    try:
        docs = await asyncio.to_thread(sql_store.get_all_documents, user_id, doc_type)
        return docs
    except Exception as e:
        raise HTTPException(500, f"문서 목록 조회 실패: {str(e)}")


@app.get("/graph/document/{doc_id}")
def graph_get_document(doc_id: str):
    """특정 문서 상세"""
    try:
        graph = get_graph_store()
        doc = graph.get_document(doc_id)
        if not doc:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_id}")
        return doc
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"문서 조회 실패: {str(e)}")


@app.delete("/graph/document/{doc_id}")
def graph_delete_document(doc_id: str):
    """Neo4j에서 문서 삭제"""
    try:
        graph = get_graph_store()
        result = graph.delete_document(doc_id)
        return {"success": True, "doc_id": doc_id, "result": result}
    except Exception as e:
        raise HTTPException(500, f"문서 삭제 실패: {str(e)}")


@app.get("/graph/document/{doc_id}/hierarchy")
def graph_get_hierarchy(doc_id: str):
    """문서 섹션 계층"""
    try:
        graph = get_graph_store()
        hierarchy = graph.get_section_hierarchy(doc_id)
        return {"doc_id": doc_id, "hierarchy": hierarchy}
    except Exception as e:
        raise HTTPException(500, f"계층 구조 조회 실패: {str(e)}")


@app.get("/graph/document/{doc_id}/references")
def graph_get_references(doc_id: str):
    """문서 참조 관계"""
    try:
        graph = get_graph_store()
        refs = graph.get_document_references(doc_id)
        if not refs:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_id}")
        return refs
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"참조 조회 실패: {str(e)}")


@app.get("/graph/search/sections")
def graph_search_sections(keyword: str, doc_id: str = None):
    """섹션 검색"""
    try:
        graph = get_graph_store()
        results = graph.search_sections(keyword, doc_id)
        return {"keyword": keyword, "results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(500, f"검색 실패: {str(e)}")


@app.get("/graph/search/terms")
def graph_search_terms(term: str):
    """용어 검색 (간소화 버전: 섹션 검색으로 대체)"""
    try:
        graph = get_graph_store()
        results = graph.search_sections(term)
        return {"term": term, "results": results, "count": len(results)}
    except Exception as e:
        raise HTTPException(500, f"용어 검색 실패: {str(e)}")


@app.get("/graph/visualization/all")
async def graph_get_full_visualization():
    """전체 문서 그래프 시각화 데이터 (모든 문서 + 관계)"""
    try:
        graph = get_graph_store()
        full_graph = await asyncio.to_thread(graph.get_full_graph)

        return {
            "success": True,
            "nodes": full_graph["nodes"],
            "links": full_graph["links"],
            "node_count": len(full_graph["nodes"]),
            "link_count": len(full_graph["links"])
        }
    except Exception as e:
        raise HTTPException(500, f"그래프 조회 실패: {str(e)}")


@app.get("/graph/visualization/{doc_id}")
def graph_get_visualization(doc_id: str, format: str = "mermaid"):
    """
    문서 관계 시각화 데이터 (프론트엔드용)

    format:
    - mermaid: Mermaid 다이어그램 코드
    - d3: D3.js용 JSON (nodes + links)
    - cytoscape: Cytoscape.js용 JSON
    """
    try:
        graph = get_graph_store()

        # 문서 참조 관계 조회
        refs = graph.get_document_references(doc_id)
        if not refs:
            raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_id}")

        doc = refs.get("document", {})
        references = refs.get("references_to", [])
        referenced_by = refs.get("referenced_by", [])

        if format == "mermaid":
            # Mermaid 다이어그램 생성
            lines = ["graph LR"]
            safe_doc_id = doc_id.replace("-", "_")
            title = doc.get("title", doc_id).replace('"', "'")

            lines.append(f'    Main[\"{doc_id}<br/>({title})\"]:::mainNode')

            # 참조하는 문서들
            for ref in references:
                ref_id = ref.get("doc_id", "").replace("-", "_")
                ref_title = ref.get("title", "").replace('"', "'")
                if ref_id:
                    lines.append(f'    Main --> {ref_id}[\"{ref.get("doc_id", "")}<br/>({ref_title})\"]')

            # 참조받는 문서들
            for cited in referenced_by:
                cited_id = cited.get("doc_id", "").replace("-", "_")
                cited_title = cited.get("title", "").replace('"', "'")
                if cited_id:
                    lines.append(f'    {cited_id}[\"{cited.get("doc_id", "")}<br/>({cited_title})\"] --> Main')

            lines.append("    classDef mainNode fill:#f96,stroke:#333,stroke-width:4px,color:#000;")
            lines.append("    classDef default fill:#eee,stroke:#333,color:#000;")

            return {
                "format": "mermaid",
                "doc_id": doc_id,
                "code": "\n".join(lines)
            }

        elif format == "d3":
            # D3.js용 JSON 생성
            nodes = []
            links = []

            # 메인 노드
            nodes.append({
                "id": doc_id,
                "label": doc.get("title", doc_id),
                "type": "main",
                "group": 0
            })

            # 참조하는 문서들 (상위)
            for idx, ref in enumerate(references):
                ref_id = ref.get("doc_id")
                if ref_id:
                    nodes.append({
                        "id": ref_id,
                        "label": ref.get("title", ref_id),
                        "type": "reference",
                        "group": 1
                    })
                    links.append({
                        "source": doc_id,
                        "target": ref_id,
                        "type": "references"
                    })

            # 참조받는 문서들 (하위)
            for idx, cited in enumerate(referenced_by):
                cited_id = cited.get("doc_id")
                if cited_id:
                    nodes.append({
                        "id": cited_id,
                        "label": cited.get("title", cited_id),
                        "type": "cited_by",
                        "group": 2
                    })
                    links.append({
                        "source": cited_id,
                        "target": doc_id,
                        "type": "cites"
                    })

            return {
                "format": "d3",
                "doc_id": doc_id,
                "data": {
                    "nodes": nodes,
                    "links": links
                }
            }

        elif format == "cytoscape":
            # Cytoscape.js용 JSON 생성
            elements = []

            # 메인 노드
            elements.append({
                "data": {
                    "id": doc_id,
                    "label": doc.get("title", doc_id),
                    "type": "main"
                },
                "classes": "main-node"
            })

            # 참조하는 문서들
            for ref in references:
                ref_id = ref.get("doc_id")
                if ref_id:
                    elements.append({
                        "data": {
                            "id": ref_id,
                            "label": ref.get("title", ref_id),
                            "type": "reference"
                        }
                    })
                    elements.append({
                        "data": {
                            "source": doc_id,
                            "target": ref_id,
                            "type": "references"
                        }
                    })

            # 참조받는 문서들
            for cited in referenced_by:
                cited_id = cited.get("doc_id")
                if cited_id:
                    elements.append({
                        "data": {
                            "id": cited_id,
                            "label": cited.get("title", cited_id),
                            "type": "cited_by"
                        }
                    })
                    elements.append({
                        "data": {
                            "source": cited_id,
                            "target": doc_id,
                            "type": "cites"
                        }
                    })

            return {
                "format": "cytoscape",
                "doc_id": doc_id,
                "elements": elements
            }

        else:
            raise HTTPException(400, f"지원하지 않는 포맷입니다: {format}")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"시각화 데이터 생성 실패: {str(e)}")


@app.get("/graph/impact/{doc_id}")
def graph_get_impact_analysis(doc_id: str):
    """
    문서 변경 시 영향 분석
    - 이 문서가 변경되면 영향을 받는 다른 문서들과 조항들
    """
    try:
        graph = get_graph_store()

        # 영향 분석 실행
        impacts = graph.get_impact_analysis(doc_id)

        if not impacts:
            return {
                "doc_id": doc_id,
                "impacts": [],
                "count": 0,
                "message": "이 문서의 변경으로 영향받는 문서가 없습니다."
            }

        # 문서별로 그룹화
        impact_by_doc = {}
        for impact in impacts:
            src_doc = impact.get("source_doc_id")
            if src_doc not in impact_by_doc:
                impact_by_doc[src_doc] = {
                    "doc_id": src_doc,
                    "sections": []
                }

            impact_by_doc[src_doc]["sections"].append({
                "section_id": impact.get("citing_section"),
                "section_title": impact.get("citing_section_title", ""),
                "context": impact.get("context", "")
            })

        return {
            "doc_id": doc_id,
            "impacts": list(impact_by_doc.values()),
            "count": len(impact_by_doc),
            "total_sections": len(impacts)
        }

    except Exception as e:
        raise HTTPException(500, f"영향 분석 실패: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
#  API 엔드포인트 - Question 추적
# ═══════════════════════════════════════════════════════════════════════════

@app.get("/graph/questions")
def graph_list_questions(limit: int = 50, session_id: str = None):
    """질문 히스토리 조회"""
    try:
        graph = get_graph_store()
        questions = graph.get_question_history(session_id=session_id, limit=limit)
        return {"questions": questions, "count": len(questions)}
    except Exception as e:
        raise HTTPException(500, f"질문 조회 실패: {str(e)}")


@app.get("/graph/questions/{question_id}/sources")
def graph_get_question_sources(question_id: str):
    """질문이 참조한 섹션 조회"""
    try:
        graph = get_graph_store()
        result = graph.get_question_sources(question_id)
        if not result:
            raise HTTPException(404, f"질문을 찾을 수 없습니다: {question_id}")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"소스 조회 실패: {str(e)}")


@app.get("/graph/stats/section-usage")
def graph_section_usage_stats(doc_id: str = None):
    """섹션 사용 통계 (간소화: Question 히스토리로 대체)"""
    try:
        graph = get_graph_store()
        # 간소화 버전: 전체 통계만 제공
        stats = graph.get_graph_stats()
        return {"stats": stats}
    except Exception as e:
        raise HTTPException(500, f"통계 조회 실패: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
#  API 엔드포인트 - 에이전트 (NEW!)
# ═══════════════════════════════════════════════════════════════════════════

# 에이전트 모듈 임포트
try:
    from backend.agent import (
        init_agent_tools, 
        run_agent, 
        AGENT_TOOLS,
        LANGCHAIN_AVAILABLE,
        LANGGRAPH_AGENT_AVAILABLE,
        ZAI_AVAILABLE
    )
    AGENT_AVAILABLE = True
    print(" 에이전트 모듈 로드 완료")
except ImportError as e:
    AGENT_AVAILABLE = False
    LANGCHAIN_AVAILABLE = False
    LANGGRAPH_AGENT_AVAILABLE = False
    ZAI_AVAILABLE = False
    print(f" 에이전트 모듈 로드 실패: {e}")


class AgentRequest(BaseModel):
    """에이전트 요청"""
    message: str
    session_id: Optional[str] = None
    llm_model: str = "gpt-4o"
    embedding_model: str = "multilingual-e5-small" # 추가
    n_results: int = DEFAULT_N_RESULTS #  추가
    use_langgraph: bool = True  # LangGraph 에이전트 사용 여부


@app.post("/agent/chat")
async def agent_chat(request: AgentRequest):
    """
     에이전트 채팅 - LLM이 도구를 선택해서 실행
    
    일반 RAG와 다르게 에이전트가 상황에 맞는 도구를 선택합니다:
    - search_sop_documents: 문서 내용 검색
    """
    if not AGENT_AVAILABLE:
        raise HTTPException(500, "에이전트 모듈이 로드되지 않았습니다")
    
    session_id = request.session_id or str(uuid.uuid4())

    # 기존 run_agent() 직접 호출 대신 큐에 넣기
    payload = request.model_dump()
    payload["session_id"] = session_id
    payload["llm_model"] = resolve_effective_llm_model(request.llm_model)

    return await enqueue_job(
        kind="agent",
        payload=payload
    )


@app.get("/agent/status")
def agent_status():
    """에이전트 상태 확인"""
    return {
        "agent_available": AGENT_AVAILABLE,
        "langchain_available": LANGCHAIN_AVAILABLE if AGENT_AVAILABLE else False,
        "langgraph_agent_available": LANGGRAPH_AGENT_AVAILABLE if AGENT_AVAILABLE else False,
        "tools": [t.name for t in AGENT_TOOLS] if AGENT_AVAILABLE else [],
        "message": "에이전트 사용 가능" if AGENT_AVAILABLE else "에이전트 모듈 로드 실패"
    }


@app.get("/agent/tools")
def agent_tools():
    """에이전트 도구 목록"""
    if not AGENT_AVAILABLE:
        raise HTTPException(500, "에이전트 모듈이 로드되지 않았습니다")

    tools_info = []
    for tool in AGENT_TOOLS:
        tools_info.append({
            "name": tool.name,
            "description": tool.description
        })

    return {"tools": tools_info, "count": len(tools_info)}


#  테스트용 간단한 에코 엔드포인트
class SimpleRequest(BaseModel):
    message: str

@app.post("/test/echo")
def test_echo(request: SimpleRequest):
    """테스트용 간단한 에코 API"""
    return {
        "session_id": str(uuid.uuid4()),
        "answer": f"테스트 응답: {request.message}",
        "success": True
    }


# ═══════════════════════════════════════════════════════════════════════════
# API 엔드포인트 - LLM as a Judge 평가
# ═══════════════════════════════════════════════════════════════════════════

class EvaluationRequest(BaseModel):
    """평가 요청 모델"""
    question: str
    answer: str
    context: Optional[str] = ""
    metrics: Optional[List[str]] = None  # ["faithfulness", "groundness", "relevancy", "correctness"]
    reference_answer: Optional[str] = None

@app.post("/evaluate")
def evaluate_answer(request: EvaluationRequest):
    """
    🔍 LLM as a Judge - 답변 평가 (RDB 검증 포함)

    평가 메트릭:
    - faithfulness: 컨텍스트 충실성 (환각 방지)
    - groundness: 근거 명확성
    - relevancy: 질문 관련성
    - correctness: 정확성과 완전성

    **무조건 RDB에서 실제 문서를 조회하여 인용 정확성 검증**
    """
    try:
        from backend.evaluation import AgentEvaluator

        # RDB 검증을 위해 sql_store 필수 전달
        evaluator = AgentEvaluator(
            judge_model="gpt-4o",
            sql_store=sql_store
        )

        # 평가 실행
        results = evaluator.evaluate_single(
            question=request.question,
            answer=request.answer,
            context=request.context,
            metrics=request.metrics,
            reference_answer=request.reference_answer
        )

        return {
            "success": True,
            "evaluation": results,
            "average_score": results.get("average_score", 0)
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"평가 실행 실패: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# === Security Context ===
# ═══════════════════════════════════════════════════════════════════════════
from passlib.context import CryptContext
from fastapi.security import OAuth2PasswordBearer
from jose import JWTError, jwt
from datetime import datetime, timedelta

pwd_context = CryptContext(schemes=["pbkdf2_sha256", "bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="auth/login")

# JWT 설정
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-change-this-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 7일

# === Auth Models ===
class Token(BaseModel):
    access_token: str
    token_type: str
    user: Dict

class TokenData(BaseModel):
    username: Optional[str] = None

class UserRegister(BaseModel):
    username: str
    password: str
    name: str
    email: str
    rank: Optional[str] = None
    dept: Optional[str] = None

    @field_validator("email")
    @classmethod
    def validate_email(cls, v):
        if not v:
            return v
        regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(regex, v):
            raise ValueError("올바른 이메일 형식이 아닙니다.")
        return v

class UserLogin(BaseModel):
    username: str
    password: str

class UserSnippet(BaseModel):
    username: str
    name: str

class PasswordReset(BaseModel):
    user_id: int
    new_password: str

class FindUsernameRequest(BaseModel):
    name: str
    dept: Optional[str] = None

# === Auth Helper Functions ===
def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)) -> Dict:
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception

    user = sql_store.get_user_by_username(token_data.username)
    if user is None:
        raise credentials_exception
    return user


# ═══════════════════════════════════════════════════════════════════════════
# === Auth Endpoints ===
# ═══════════════════════════════════════════════════════════════════════════

# 부서/직책 기본 목록
DEFAULT_DEPARTMENTS = ["품질관리부", "품질보증부", "생산부", "연구개발부", "경영지원부", "영업부"]
DEFAULT_RANKS = ["사원", "주임", "대리", "과장", "차장", "부장", "이사", "상무"]

@app.get("/auth/options")
def get_auth_options():
    """회원가입 시 부서/직책 드롭다운 옵션 조회"""
    depts = set(DEFAULT_DEPARTMENTS)
    ranks = set(DEFAULT_RANKS)
    try:
        with sql_store._get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT DISTINCT dept FROM users WHERE dept IS NOT NULL AND dept != ''")
                for row in cur.fetchall():
                    depts.add(row[0])
                cur.execute("SELECT DISTINCT rank FROM users WHERE rank IS NOT NULL AND rank != ''")
                for row in cur.fetchall():
                    ranks.add(row[0])
    except Exception:
        pass
    return {
        "departments": sorted(list(depts)),
        "ranks": sorted(list(ranks))
    }

@app.post("/auth/register", response_model=Token)
async def register(user: UserRegister):
    # 중복 체크
    existing_user = sql_store.get_user_by_username(user.username)
    if existing_user:
        raise HTTPException(status_code=400, detail="이미 존재하는 아이디입니다.")
    
    # 비밀번호 해시 & 저장
    hashed_pw = get_password_hash(user.password)
    user_id = sql_store.register_user(
        username=user.username,
        password_hash=hashed_pw,
        name=user.name,
        email=user.email,
        rank=user.rank,
        dept=user.dept
    )
    
    if not user_id:
        raise HTTPException(status_code=500, detail="회원가입 처리에 실패했습니다.")
    
    # 회원가입 직후 로그인 상태로 간주하여 last_login 갱신
    sql_store.update_last_login(user_id)
    
    # 자동 로그인 처리 (토큰 발급)
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username, "user_id": user_id},
        expires_delta=access_token_expires
    )
    
    # 사용자 정보 조회 (응답용)
    new_user = sql_store.get_user(user_id)
    
    return {"access_token": access_token, "token_type": "bearer", "user": new_user}

@app.post("/auth/login", response_model=Token)
async def login(user_req: UserLogin):
    user = sql_store.get_user_by_username(user_req.username)
    if not user or not verify_password(user_req.password, user['password_hash']):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="아이디 또는 비밀번호가 올바르지 않습니다.",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    # 로그인 성공: last_login 갱신 & 토큰 발급
    sql_store.update_last_login(user['id'])
    # 갱신된 정보(last_login 등)를 다시 가져옴
    user = sql_store.get_user_by_username(user_req.username)
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user['username'], "user_id": user['id']},
        expires_delta=access_token_expires
    )
    
    # 민감 정보 제거
    user_resp = {k: v for k, v in user.items() if k != 'password_hash'}
    
    return {"access_token": access_token, "token_type": "bearer", "user": user_resp}

@app.get("/auth/me")
async def read_users_me(current_user: Dict = Depends(get_current_user)):
    user_resp = {k: v for k, v in current_user.items() if k != 'password_hash'}
    return {"user": user_resp}

@app.post("/auth/find-username")
async def find_username(req: FindUsernameRequest):
    """이름과 부서로 아이디 찾기"""
    with sql_store._get_connection() as conn:
        with conn.cursor() as cur:
            query = "SELECT username FROM users WHERE name = %s"
            params = [req.name]
            if req.dept and req.dept != '전체':
                query += " AND dept = %s"
                params.append(req.dept)
            
            cur.execute(query, tuple(params))
            res = cur.fetchone()
            if res:
                return {"username": res[0]}
            raise HTTPException(status_code=404, detail="일치하는 사용자를 찾을 수 없습니다.")

@app.post("/auth/verify-user")
async def verify_user_identity(req: UserSnippet):
    """비밀번호 재설정을 위한 본인 확인 (아이디 + 이름)"""
    user = sql_store.get_user_by_username(req.username)
    if user and user['name'] == req.name:
        return {"user_id": user['id']}
    raise HTTPException(status_code=404, detail="사용자 정보가 일치하지 않습니다.")

@app.post("/auth/reset-password")
async def reset_password_endpoint(req: PasswordReset):
    """비밀번호 재설정"""
    hashed_pw = get_password_hash(req.new_password)
    try:
        with sql_store._get_connection() as conn:
            with conn.cursor() as cur:
                cur.execute("UPDATE users SET password_hash = %s WHERE id = %s", (hashed_pw, req.user_id))
                conn.commit()
        return {"message": "비밀번호가 변경되었습니다."}
    except Exception:
        raise HTTPException(status_code=500, detail="비밀번호 변경 실패")


# ═══════════════════════════════════════════════════════════════════════════
# OnlyOffice + S3 엔드포인트
# ═══════════════════════════════════════════════════════════════════════════

class OnlyOfficeConfigRequest(BaseModel):
    doc_name: str
    version: Optional[str] = None
    user_name: str = "편집자"
    mode: str = "view"


async def process_docx_upload_task(
    content: bytes,
    doc_name: str,
    version: str,
    collection: str,
    task_id: str
):
    """DOCX 업로드 및 RAG 파이프라인 백그라운드 작업"""
    start_time = time.time()
    update_task_status(task_id, "processing", f"'{doc_name}' DOCX 파일을 분석 중입니다.")
    
    try:
        # 1. S3 저장
        s3 = get_s3_store()
        s3_key = await asyncio.to_thread(s3.upload_docx, doc_name, version, content)
        
        # 2. 파이프라인 실행
        model_path = resolve_model_path("multilingual-e5-small")
        embed_model = SentenceTransformer(model_path)
        filename = f"{doc_name}_v{version}.docx"

        result = await asyncio.to_thread(
            process_document,
            file_path=filename,
            content=content,
            doc_id=doc_name,
            use_llm_metadata=False,
            embed_model=embed_model,
        )

        if not result.get("success"):
            raise Exception(f"문서 처리 실패: {result.get('errors')}")

        chunks_data = result.get("chunks", [])
        
        @dataclass
        class _Chunk:
            text: str
            metadata: dict
            index: int = 0

        chunks = [_Chunk(text=c["text"], metadata=c["metadata"], index=c["index"]) for c in chunks_data]
        update_task_status(task_id, "processing", f"DOCX 파싱 완료 (청크 {len(chunks)}개). DB 저장 중...")

        # 3. Weaviate 저장
        pipeline_version = "pdf-clause-v2.0"
        texts = [c.text for c in chunks]
        metadatas = [
            {**c.metadata, "chunk_method": "article", "model": "multilingual-e5-small", "pipeline_version": pipeline_version}
            for c in chunks
        ]
        await asyncio.to_thread(vector_store.add_documents, texts=texts, metadatas=metadatas, collection_name=collection, model_name=model_path)

        # 4. PostgreSQL 저장
        from backend.document_pipeline import docx_to_markdown
        markdown_text = await asyncio.to_thread(docx_to_markdown, content)

        doc_id_db = await asyncio.to_thread(
            sql_store.save_document,
            doc_name=doc_name,
            content=markdown_text,
            doc_type="docx",
            version=version,
        )
        if doc_id_db and chunks:
            batch_chunks = [
                {"clause": c.metadata.get("clause_id"), "content": c.text, "metadata": c.metadata}
                for c in chunks
            ]
            await asyncio.to_thread(sql_store.save_chunks_batch, doc_id_db, batch_chunks)

        # 5. Neo4j 저장
        try:
            graph = get_graph_store()
            if graph and graph.test_connection():
                await asyncio.to_thread(_upload_to_neo4j_from_pipeline, graph, result, filename)
        except Exception as graph_err:
            print(f"  Neo4j 저장 실패 (건너뜀): {graph_err}")

        elapsed = round(time.time() - start_time, 2)
        update_task_status(task_id, "completed", f"DOCX 업로드 및 분석이 완료되었습니다. ({elapsed}초)", doc_name=doc_name, version=version, s3_key=s3_key)

    except Exception as e:
        import traceback
        traceback.print_exc()
        update_task_status(task_id, "error", f"DOCX 처리 중 오류 발생: {str(e)}")


@app.post("/rag/upload-docx")
async def upload_docx_to_s3(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    doc_name: str = Form(...),
    version: str = Form("1.0"),
    collection: str = Form("documents"),
):
    """
    DOCX 파일을 S3에 저장 후 RAG 파이프라인 실행 (비동기)
    """
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(400, "DOCX 파일만 업로드 가능합니다.")

    content = await file.read()
    task_id = f"docx_{uuid.uuid4().hex[:8]}"
    update_task_status(task_id, "waiting", f"'{doc_name}' DOCX 업로드 요청이 접수되었습니다.", doc_name=doc_name)

    background_tasks.add_task(
        process_docx_upload_task,
        content=content,
        doc_name=doc_name,
        version=version,
        collection=collection,
        task_id=task_id
    )

    return {
        "success": True,
        "message": "DOCX 업로드 및 분석이 시작되었습니다.",
        "task_id": task_id,
        "doc_name": doc_name
    }


class UploadS3Request(BaseModel):
    s3_key: str
    doc_name: str
    version: str = "1.0"


# S3 / OnlyOffice 싱글톤 (필요 시 lazy 초기화)
_s3_store = None

def get_s3_store():
    global _s3_store
    if _s3_store is None:
        from backend.s3_store import S3Store
        _s3_store = S3Store()
    return _s3_store


@app.post("/onlyoffice/config")
async def onlyoffice_config(request: OnlyOfficeConfigRequest):
    """
    OnlyOffice 에디터 설정 JSON 반환

    입력: { doc_name, version(optional), user_name }
    동작:
      1. SQL에서 최신 버전 조회 (version 미지정 시)
      2. S3 presigned URL 생성
      3. OnlyOffice 설정 JSON 반환
    """
    try:
        from backend.onlyoffice_service import create_editor_config, get_onlyoffice_server_url, BACKEND_URL

        doc_name = request.doc_name

        # 버전 결정: 사용자 지정 없으면 DB 최신 버전 사용
        version = request.version
        if not version:
            versions_data = sql_store.get_document_versions(doc_name)
            if not versions_data:
                raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name}")
            version = versions_data[0].get('version', '1.0')

        # 백엔드 내부 URL 사용 (S3 presigned URL 대신)
        # OnlyOffice가 JWT 헤더를 붙여 요청 → 백엔드가 검증 후 S3에서 파일 서빙
        file_url = f"{BACKEND_URL}/onlyoffice/document/{doc_name}/{version}"

        # OnlyOffice 설정 생성
        config = create_editor_config(
            doc_id=doc_name,
            version=version,
            user_name=request.user_name,
            file_url=file_url,
            mode=request.mode,
        )

        return {
            "config": config,
            "doc_name": doc_name,
            "version": version,
            "onlyoffice_server_url": get_onlyoffice_server_url(),
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"OnlyOffice 설정 생성 실패: {str(e)}")


@app.get("/onlyoffice/document/{doc_name}/{version}")
async def serve_docx_for_onlyoffice(doc_name: str, version: str):
    """
    OnlyOffice Document Server가 DOCX를 가져가는 내부 엔드포인트.
    S3 presigned URL 대신 이 URL을 document.url로 사용:
      - OnlyOffice가 JWT Authorization 헤더를 붙여도 S3와 충돌 없음
      - 백엔드가 S3에서 파일을 가져와 OnlyOffice에 직접 서빙
    """
    try:
        s3 = get_s3_store()
        content = s3.download_docx(doc_name, version)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{doc_name}_v{version}.docx"'},
        )
    except Exception as e:
        raise HTTPException(404, f"문서를 찾을 수 없습니다: {doc_name} v{version} ({e})")


@app.post("/onlyoffice/callback")
async def onlyoffice_callback(request: Request):
    """
    OnlyOffice 콜백 처리

    OnlyOffice가 저장 완료 시 호출.
    status 2(저장 중) 또는 6(편집 완료) 시 DOCX를 S3에 새 버전으로 저장하고
    RAG 파이프라인을 실행합니다.

    반환: {"error": 0} (OnlyOffice 요구사항)
    """
    try:
        callback_data = await request.json()
        status = callback_data.get('status')
        download_url = callback_data.get('url')
        key = callback_data.get('key', '')

        print(f"\n[OnlyOffice Callback] status={status}, key={key}")

        # status 2 = 문서 저장 중, status 6 = 편집 완료(강제 저장)
        if status not in (2, 6):
            return {"error": 0}

        if not download_url:
            print("  콜백 URL 없음 - 건너뜀")
            return {"error": 0}

        # key 형식: {doc_id}_v{version}_{timestamp}
        parts = key.split('_v')
        doc_id = parts[0] if len(parts) > 1 else key
        
        task_id = f"onlyoffice_{doc_id}_{int(time.time())}"
        update_task_status(task_id, "processing", f"'{doc_id}' OnlyOffice 편집본을 저장 중입니다.", doc_name=doc_id)
        version_part = parts[1].rsplit('_', 1)[0] if len(parts) > 1 else '1.0'

        # 새 버전 번호 결정 (현재 버전 + 0.1)
        try:
            current_v = float(version_part)
            new_version = f"{current_v + 0.1:.1f}"
        except ValueError:
            new_version = version_part + "_edited"

        print(f"  문서: {doc_id}, 현재버전: {version_part} → 새버전: {new_version}")

        # 1. OnlyOffice 서버에서 편집된 DOCX 다운로드
        from backend.onlyoffice_service import download_from_onlyoffice
        docx_content = await download_from_onlyoffice(download_url)
        print(f"  DOCX 다운로드 완료: {len(docx_content)} bytes")

        # 2. S3에 새 버전으로 저장
        s3 = get_s3_store()
        s3_key = s3.upload_docx(doc_id, new_version, docx_content)
        print(f"  S3 저장 완료: {s3_key}")

        # 3. DOCX → 마크다운 변환
        from backend.document_pipeline import docx_to_markdown
        markdown_text = docx_to_markdown(docx_content)

        # 4. RAG 파이프라인 실행 (process_document)
        model_path = resolve_model_path("multilingual-e5-small")
        embed_model = SentenceTransformer(model_path)

        result = process_document(
            file_path=f"{doc_id}_v{new_version}.docx",
            content=docx_content,
            doc_id=doc_id,
            use_llm_metadata=False,
            embed_model=embed_model,
        )

        if not result.get("success"):
            print(f"  파이프라인 실패: {result.get('errors')}")
            return {"error": 0}

        chunks_data = result.get("chunks", [])

        # 5. Weaviate 저장
        from dataclasses import dataclass

        @dataclass
        class _Chunk:
            text: str
            metadata: dict
            index: int = 0

        chunks = [_Chunk(text=c["text"], metadata=c["metadata"], index=c["index"]) for c in chunks_data]
        texts = [c.text for c in chunks]
        metadatas = [
            {**c.metadata, "chunk_method": "article", "model": "multilingual-e5-small", "pipeline_version": "pdf-clause-v2.0"}
            for c in chunks
        ]
        vector_store.add_documents(
            texts=texts, metadatas=metadatas, collection_name="documents", model_name=model_path
        )

        # 6. PostgreSQL 저장
        doc_id_db = sql_store.save_document(
            doc_name=doc_id,
            content=markdown_text,
            doc_type="docx",
            version=new_version,
        )
        if doc_id_db and chunks:
            batch_chunks = [
                {"clause": c.metadata.get("clause_id"), "content": c.text, "metadata": c.metadata}
                for c in chunks
            ]
            sql_store.save_chunks_batch(doc_id_db, batch_chunks)

        # 7. Neo4j 저장
        try:
            graph = get_graph_store()
            if graph.test_connection():
                _upload_to_neo4j_from_pipeline(graph, result, f"{doc_id}_v{new_version}.docx")
        except Exception as graph_err:
            print(f"  Neo4j 저장 실패 (건너뜀): {graph_err}")

        elapsed = round(time.time() - start_time, 2)
        update_task_status(task_id, "completed", f"문서 수정이 완료되었습니다. ({elapsed}초)", doc_id=doc_id, version=new_version)
        print(f"  [OnlyOffice Callback] 완료 - 새 버전 {new_version} 저장됨")
        return {"error": 0}

    except Exception as e:
        import traceback
        traceback.print_exc()
        if 'task_id' in locals():
            update_task_status(task_id, "error", f"OnlyOffice 저장 실패: {str(e)}")
        print(f"  [OnlyOffice Callback] 오류: {e}")
        # OnlyOffice는 error: 0이 아니면 재시도하므로 항상 0 반환
        return {"error": 0}


@app.post("/rag/upload-s3")
async def upload_from_s3(request: UploadS3Request):
    """
    S3에 이미 있는 DOCX 파일을 RAG 파이프라인으로 처리

    입력: { s3_key, doc_name, version }
    동작:
      1. S3에서 DOCX 다운로드
      2. 기존 /rag/upload 파이프라인 실행
    """
    start_time = time.time()
    try:
        import boto3
        import os

        # S3에서 파일 다운로드
        s3_client = boto3.client(
            's3',
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=os.getenv('AWS_REGION', 'ap-northeast-2'),
        )
        bucket = os.getenv('S3_BUCKET_NAME')
        response = s3_client.get_object(Bucket=bucket, Key=request.s3_key)
        content = response['Body'].read()

        filename = f"{request.doc_name}_v{request.version}.docx"
        model_path = resolve_model_path("multilingual-e5-small")
        embed_model = SentenceTransformer(model_path)

        result = process_document(
            file_path=filename,
            content=content,
            doc_id=request.doc_name,
            use_llm_metadata=False,
            embed_model=embed_model,
        )

        if not result.get("success"):
            raise HTTPException(400, f"문서 처리 실패: {result.get('errors')}")

        chunks_data = result.get("chunks", [])

        from dataclasses import dataclass

        @dataclass
        class _Chunk:
            text: str
            metadata: dict
            index: int = 0

        chunks = [_Chunk(text=c["text"], metadata=c["metadata"], index=c["index"]) for c in chunks_data]

        # Weaviate 저장
        pipeline_version = "pdf-clause-v2.0"
        texts = [c.text for c in chunks]
        metadatas = [
            {**c.metadata, "chunk_method": "article", "model": "multilingual-e5-small", "pipeline_version": pipeline_version}
            for c in chunks
        ]
        vector_store.add_documents(
            texts=texts, metadatas=metadatas, collection_name="documents", model_name=model_path
        )

        # PostgreSQL 저장
        from backend.document_pipeline import docx_to_markdown
        markdown_text = docx_to_markdown(content)

        final_version = request.version or result.get("version", "1.0")
        doc_id_db = sql_store.save_document(
            doc_name=request.doc_name,
            content=markdown_text,
            doc_type="docx",
            version=final_version,
        )
        if doc_id_db and chunks:
            batch_chunks = [
                {"clause": c.metadata.get("clause_id"), "content": c.text, "metadata": c.metadata}
                for c in chunks
            ]
            sql_store.save_chunks_batch(doc_id_db, batch_chunks)

        # Neo4j 저장
        try:
            graph = get_graph_store()
            if graph.test_connection():
                _upload_to_neo4j_from_pipeline(graph, result, filename)
        except Exception as graph_err:
            print(f"  Neo4j 저장 실패 (건너뜀): {graph_err}")

        elapsed = round(time.time() - start_time, 2)
        return {
            "success": True,
            "doc_name": request.doc_name,
            "version": final_version,
            "chunks": len(chunks),
            "elapsed_seconds": elapsed,
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"S3 업로드 처리 실패: {str(e)}")


# ═══════════════════════════════════════════════════════════════════════════
# 서버 실행
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print("[시스템] 초기화 중...")
    sql_store.init_db()
    
    # Neo4j 연결 확인 (성공 로그는 connect 내부에서 출력됨)
    try:
        get_graph_store()
    except Exception as e:
        print(f" Neo4j 초기 연결 실패: {e}")

    # Weaviate 연결 확인 (성공 로그는 get_client 내부에서 출력됨)
    try:
        wv_client = vector_store.get_client()
        if not wv_client.is_connected():
            print(" Weaviate v4 연결 상태 확인 실패")
    except Exception as e:
        print(f" Weaviate v4 연결 체크 중 오류: {e}")

    
    import uvicorn
    
    print("\n" + "=" * 60)
    print(" RAG Chatbot API v14.0 + OpenAI Agent")
    print("=" * 60)
    print(f" LLM 백엔드: OpenAI (gpt-4o)")
    print(f" 에이전트: {' 활성화' if LANGGRAPH_AVAILABLE else ' 비활성화'}")
    
    if LANGGRAPH_AVAILABLE:
        print(f"   - LangChain: {'활성화' if LANGCHAIN_AVAILABLE else '비활성화'}")
    print("Docs: http://localhost:8000/docs")
    print("=" * 60)
    print("주요 기능:")
    print("  - LangGraph 문서 파이프라인")
    print("  -  ReAct 에이전트 (/agent/chat)")
    print("  - Weaviate(v4) + Neo4j + PostgreSQL")
    print("  - LangSmith 추적 지원")
    print("=" * 60)
    
    uvicorn.run(app, host="0.0.0.0", port=8000, workers=1)


if __name__ == "__main__":
    main()
