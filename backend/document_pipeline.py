"""
LangGraph ê¸°ë°˜ ë¬¸ì„œ ì²˜ë¦¬ íŒŒì´í”„ë¼ì¸ v9.1

ğŸ”¥ v9.1 ê°œì„ :
- í˜ì´ì§€ ë²ˆí˜¸ ë©”íƒ€ë°ì´í„° ì¶”ê°€
- Parent-Child ê³„ì¸µ êµ¬ì¡° ë„ì…
- ë¬¸ì„œ í—¤ë” ë©”íƒ€ë°ì´í„° ì¶”ì¶œ (SOP ID, Version, Date, Department)
- ëª©ì°¨ ì¤‘ë³µ ì œê±°
- ë©”íƒ€ë°ì´í„° ìµœì í™” (ë¶ˆí•„ìš” í•„ë“œ ì œê±°)

ë…¸ë“œ íë¦„:
â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”
â”‚  Load   â”‚â”€â”€â”€â–¶â”‚ Convert â”‚â”€â”€â”€â–¶â”‚ Validate â”‚â”€â”€â”€â–¶â”‚  Split  â”‚
â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜    â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜    â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜    â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜
                    â”‚              â”‚                â”‚
                    â–¼              â–¼                â–¼
              â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”  â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”    â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
              â”‚ Fallback â”‚  â”‚  Repair  â”‚    â”‚ Optimize â”‚
              â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜  â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜    â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜
                                                   â”‚
                                                   â–¼
                                            â”Œâ”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”
                                            â”‚ Finalize â”‚
                                            â””â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”˜
"""

from typing import TypedDict, List, Dict, Optional, Literal, Annotated
from dataclasses import dataclass, field
import re
from io import BytesIO
import operator
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from .llm import get_llm_response


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ìƒíƒœ ì •ì˜
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

class PipelineState(TypedDict):
    """íŒŒì´í”„ë¼ì¸ ìƒíƒœ"""
    # ì…ë ¥
    filename: str
    content: bytes

    # ì„¤ì •
    chunk_size: int
    chunk_overlap: int
    use_llm_metadata: bool  # LLM ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ì‚¬ìš© ì—¬ë¶€
    use_clause_parsing: bool  # ì¡°í•­ ë²ˆí˜¸ ê¸°ë°˜ íŒŒì‹± ì‚¬ìš© ì—¬ë¶€

    # ì¤‘ê°„ ê²°ê³¼
    file_type: str
    markdown: str
    metadata: Dict
    sections: List[Dict]
    chunks: List[Dict]

    # í’ˆì§ˆ ì§€í‘œ
    quality_score: float
    conversion_method: str

    # ì—ëŸ¬ ì²˜ë¦¬
    errors: Annotated[List[str], operator.add]
    warnings: Annotated[List[str], operator.add]
    retry_count: int

    # ìµœì¢… ê²°ê³¼
    success: bool


@dataclass
class Chunk:
    """ì²­í¬ ë°ì´í„°"""
    text: str
    index: int = 0
    metadata: Dict = field(default_factory=dict)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ë¬¸ì„œ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def extract_document_metadata(text: str, filename: str) -> Dict:
    """
    ë¬¸ì„œ í—¤ë”ì—ì„œ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ (Pure AI: LLM Only)
    """
    print(f"ğŸ§  [Metadata] AI ê¸°ë°˜ ì§€ëŠ¥í˜• ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ì¤‘... (File: {filename})")
    
    head_text = text[:3000]
    
    prompt = f"""ë‹¹ì‹ ì€ GMP ê·œì •(SOP) ë¶„ì„ ì „ë¬¸ê°€ì…ë‹ˆë‹¤. ë‹¤ìŒ ë¬¸ì„œì˜ ë‚´ìš©ì„ ë¶„ì„í•˜ì—¬ ê´€ë¦¬ìš© ë©”íƒ€ë°ì´í„°ë¥¼ ì¶”ì¶œí•˜ì„¸ìš”.

[ì¶”ì¶œ ê·œì¹™]
1. doc_id: 'EQ-SOP-0001'ê³¼ ê°™ì€ ê´€ë¦¬ ë²ˆí˜¸ë¥¼ ì°¾ìœ¼ì„¸ìš”.
2. title: ë¬¸ì„œì˜ ê³µì‹ ëª…ì¹­ì„ ì •í™•íˆ ì¶”ì¶œí•˜ì„¸ìš”. 
3. version: '1.0' ë˜ëŠ” 'Ver 2.1' ê°™ì€ í˜•ì‹ì„ ì°¾ìœ¼ì„¸ìš”.
4. effective_date: YYYY-MM-DD í˜•ì‹ìœ¼ë¡œ ë³€í™˜í•˜ì„¸ìš”.
5. owning_dept: ìƒì‚°íŒ€, í’ˆì§ˆë³´ì¦íŒ€ ë“± ë‹´ë‹¹ ì¡°ì§ëª…ì„ ì°¾ìœ¼ì„¸ìš”.

ê²°ê³¼ëŠ” ë°˜ë“œì‹œ ì•„ë˜ JSON í˜•ì‹ìœ¼ë¡œë§Œ ë‹µë³€í•˜ì„¸ìš”.
{{
  "doc_id": "ì¶”ì¶œëœ ID (ì—†ìœ¼ë©´ null)",
  "title": "ë¬¸ì„œ ì œëª© (í•„ìˆ˜)",
  "version": "ì¶”ì¶œëœ ë²„ì „ (ì—†ìœ¼ë©´ null)",
  "effective_date": "YYYY-MM-DD (ì—†ìœ¼ë©´ null)",
  "owning_dept": "ë‹´ë‹¹ ë¶€ì„œ (ì—†ìœ¼ë©´ null)"
}}

[íŒŒì¼ëª…]
{filename}

[ì£¼ì˜ì‚¬í•­]
ë°˜ë“œì‹œ ìƒê° ê³¼ì •(Reasoning)ì„ ìƒëµí•˜ê³ , ìµœì¢… JSON ë°ì´í„°ë§Œ ì¶œë ¥í•˜ì„¸ìš”.

[ë¬¸ì„œ ë‚´ìš©]
{head_text}"""
    
    metadata = {"file_name": filename}
    try:
        # GPT-4o-mini: JSON ì¶”ì¶œì— ìµœì í™”
        llm_res = get_llm_response(prompt, llm_model="gpt-4o-mini", max_tokens=4096, temperature=0.1)
        json_match = re.search(r'\{.*\}', llm_res, re.DOTALL)
        if json_match:
            llm_meta = json.loads(json_match.group(0))
            # ğŸ”¥ í˜¸í™˜ì„± ë ˆì´ì–´: doc_idë¥¼ sop_idë¡œë„ ë³µì‚¬
            if 'doc_id' in llm_meta and 'sop_id' not in llm_meta:
                llm_meta['sop_id'] = llm_meta['doc_id']
            metadata.update(llm_meta)
            print(f"âœ… [Metadata] AI ì¶”ì¶œ ì„±ê³µ: {metadata.get('doc_id') or metadata.get('sop_id') or 'ID ë¯¸í™•ì¸'}")
    except Exception as e:
        print(f"âš ï¸ [Metadata] AI ì¶”ì¶œ ì‹¤íŒ¨: {e}")
        metadata.update({"doc_id": None, "sop_id": None, "title": filename, "version": None, "effective_date": None, "owning_dept": None})

    return metadata

def extract_clause_metadata(text: str, doc_info: Dict, section_name: str) -> Dict:
    """
    ì¡°í•­(Clause) ë‹¨ìœ„ ìƒì„¸ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ (processor ë°©ì‹)
    """
    # ğŸ”¥ ë„ˆë¬´ ì§§ì€ ë‚´ìš©ì€ ë¶„ì„ ìŠ¤í‚µ
    clean_text = text.strip()
    if len(clean_text) < 30:
        return {}

    # print(f"ğŸ§  [Clause Scan] {section_name} ìƒì„¸ ë¶„ì„ ì¤‘...")

    prompt = f"""ë‹¹ì‹ ì€ GMP ë¬¸ì„œ ë¶„ì„ ì „ë¬¸ê°€ì…ë‹ˆë‹¤.
ë‹¤ìŒ ë¬¸ì„œ ì²­í¬ë¥¼ ë¶„ì„í•´ì„œ ê²€ìƒ‰ì— ìœ ìš©í•œ ë©”íƒ€ë°ì´í„°ë¥¼ ì¶”ì¶œí•˜ì„¸ìš”.

## ì²­í¬ ì •ë³´
- ì¡°í•­ë²ˆí˜¸: {section_name}
- ì œëª©: {section_name}
- ë‚´ìš©:
{text[:1000]}

## ì¶”ì¶œí•  ë©”íƒ€ë°ì´í„°
ë‹¤ìŒ í•­ëª©ë“¤ì„ JSON í˜•ì‹ìœ¼ë¡œ ì¶”ì¶œí•˜ì„¸ìš”:

1. content_type: ì´ ì²­í¬ê°€ ì„¤ëª…í•˜ëŠ” ë‚´ìš©ì˜ ìœ í˜• (ì˜ˆ: ëª©ì , ì •ì˜, ì±…ì„, ì ˆì°¨, ê¸°ì¤€, ê¸°ë¡, ì°¸ê³ ë¬¸í—Œ ë“±)
2. main_topic: ì´ ì²­í¬ì˜ í•µì‹¬ ì£¼ì œ
3. sub_topics: ê´€ë ¨ ì„¸ë¶€ ì£¼ì œë“¤ (ë¦¬ìŠ¤íŠ¸, ìµœëŒ€ 3ê°œ)
4. actors: ì–¸ê¸‰ëœ ì—­í• ì/ë‹´ë‹¹ì (ë¦¬ìŠ¤íŠ¸)
5. actions: ìˆ˜í–‰í•´ì•¼ í•˜ëŠ” í–‰ìœ„/ì ˆì°¨ (ë¦¬ìŠ¤íŠ¸, ìµœëŒ€ 3ê°œ)
6. conditions: íŠ¹ìˆ˜ ì¡°ê±´ì´ë‚˜ ìƒí™© (ë¦¬ìŠ¤íŠ¸)
7. summary: í•œ ë¬¸ì¥ ìš”ì•½ (30ì ì´ë‚´)
8. intent_scope: ì´ ì¡°í•­ì´ ë‹¤ë£¨ëŠ” ê´€ë¦¬ ì˜ì—­ (ë‹¤ìŒ ì¤‘ 1ê°œë§Œ ì„ íƒ)
   - user_account: ì‚¬ìš©ì ê³„ì •Â·ê¶Œí•œÂ·ì—­í•  ê´€ë¦¬
   - document_lifecycle: ë¬¸ì„œì˜ ìˆ˜ëª…ì£¼ê¸° (ì‘ì„±, ìŠ¹ì¸, ê°œì •, íê¸° ë“±)
   - system_configuration: ì‹œìŠ¤í…œ ì„¤ì • or êµ¬ì¡° ë³€ê²½
   - audit_evidence: ê°ì‚¬ëŒ€ì‘ì— ê´€ë ¨í•œ ìë£Œ
   - training: êµìœ¡, í›ˆë ¨, ìê²©
9. intent_summary: ì´ ì¡°í•­ì´ ì–´ë–¤ ì§ˆë¬¸ì— ë‹µí•˜ëŠ”ì§€ë¥¼ ì˜ì–´ 1ë¬¸ì¥ìœ¼ë¡œ ìš”ì•½ (ì˜ˆ: "What is the purpose of Level 1 quality manual?")
10. language: ì´ ì²­í¬ì˜ ì£¼ìš” ì–¸ì–´ ("ko" ë˜ëŠ” "en")

## ì¶œë ¥
JSONë§Œ ì¶œë ¥:
{{"content_type": "...", "main_topic": "...", "sub_topics": [...], "actors": [...], "actions": [...], "conditions": [...], "summary": "...", "intent_scope": "...", "intent_summary": "...", "language": "..."}}"""

    try:
        # GPT-4o-mini: ë¹ ë¥´ê³  ì •í™•í•œ ì¡°í•­ë³„ ë©”íƒ€ë°ì´í„° JSON ì¶”ì¶œ
        llm_res = get_llm_response(prompt, llm_model="gpt-4o-mini", max_tokens=4096, temperature=0.1)
        # ğŸ”¥ processor ë°©ì‹: JSON íŒŒì‹±
        result_text = llm_res.strip()

        # JSON ì½”ë“œ ë¸”ë¡ ì²˜ë¦¬
        if "```" in result_text:
            json_match = result_text.split("```")[1]
            if json_match.startswith("json"):
                json_match = json_match[4:]
            result_text = json_match.strip()

        res = json.loads(result_text)

        # ë¦¬ìŠ¤íŠ¸ë¥¼ ì‰¼í‘œ êµ¬ë¶„ ë¬¸ìì—´ë¡œ ë³€í™˜ (ê¸°ì¡´ pipeline ë°©ì‹ê³¼ í˜¸í™˜)
        if isinstance(res.get('sub_topics'), list):
            res['sub_topics'] = ', '.join(res['sub_topics'])
        if isinstance(res.get('actors'), list):
            res['actors'] = ', '.join(res['actors'])
        if isinstance(res.get('actions'), list):
            res['actions'] = ', '.join(res['actions'])
        if isinstance(res.get('conditions'), list):
            res['conditions'] = ', '.join(res['conditions'])

        # í˜¸í™˜ì„± ë³´ì¥
        if 'doc_id' not in res and doc_info.get('doc_id'):
            res['doc_id'] = doc_info.get('doc_id')

        return res
    except json.JSONDecodeError as e:
        print(f"âš ï¸ [Clause Scan] JSON íŒŒì‹± ì‹¤íŒ¨: {e}")
        return {
            "content_type": "",
            "main_topic": "",
            "sub_topics": "",
            "actors": "",
            "actions": "",
            "conditions": "",
            "summary": "",
            "intent_scope": "",
            "intent_summary": "",
            "language": "ko"
        }
    except Exception as e:
        print(f"âš ï¸ [Clause Scan] ì‹¤íŒ¨: {e}")
        return {}

    return {}


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ë…¸ë“œ í•¨ìˆ˜ë“¤
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def node_load(state: PipelineState) -> PipelineState:
    """
    1ë‹¨ê³„: íŒŒì¼ ë¡œë“œ ë° íƒ€ì… ê°ì§€
    """
    filename = state["filename"]
    content = state["content"]
    
    # ğŸ”¥ ì‹¤ì œ í™•ì¥ì ì¶”ì¶œ (ë§ˆì§€ë§‰ . ì´í›„)
    filename_lower = filename.lower()
    if '.' in filename_lower:
        actual_ext = filename_lower.rsplit('.', 1)[-1]
    else:
        actual_ext = ''
    
    # í™•ì¥ì ê¸°ë°˜ íƒ€ì… ê²°ì •
    if actual_ext == 'pdf':
        file_type = 'pdf'
    elif actual_ext == 'docx':
        file_type = 'docx'
    elif actual_ext == 'doc':
        file_type = 'doc'
    elif actual_ext in ['html', 'htm']:
        file_type = 'html'
    elif actual_ext == 'md':
        file_type = 'markdown'
    elif actual_ext == 'txt':
        file_type = 'text'
    else:
        file_type = 'unknown'
    
    # íŒŒì¼ í¬ê¸° í™•ì¸
    file_size = len(content)
    if file_size == 0:
        state["errors"] = ["íŒŒì¼ì´ ë¹„ì–´ìˆìŠµë‹ˆë‹¤."]
        state["success"] = False
        return state
    
    state["file_type"] = file_type
    state["metadata"] = {
        "file_name": filename,
        "file_type": file_type,
        "file_size": file_size,
    }
    
    return state


def node_convert(state: PipelineState) -> PipelineState:
    """
    2ë‹¨ê³„: ë¬¸ì„œ â†’ ë§ˆí¬ë‹¤ìš´ ë³€í™˜
    """
    file_type = state["file_type"]
    content = state["content"]
    filename = state["filename"]
    
    try:
        if file_type == 'docx':
            markdown, metadata = _convert_docx(filename, content)
            state["conversion_method"] = "python-docx"
            
        elif file_type == 'pdf':
            markdown, metadata, method = _convert_pdf_with_fallback(filename, content)
            state["conversion_method"] = method
            # ğŸ”¥ PDFëŠ” í—¤ë” ì¶”ë¡  í•„ìˆ˜!
            markdown = _infer_headers(markdown)
            state["conversion_method"] += "+infer-headers"
            
        elif file_type == 'html':
            markdown, metadata = _convert_html(filename, content)
            state["conversion_method"] = "beautifulsoup"
            
        elif file_type == 'markdown':
            markdown = content.decode('utf-8', errors='ignore')
            metadata = {}
            state["conversion_method"] = "passthrough"
            
        elif file_type == 'text':
            markdown = _convert_text_to_markdown(content.decode('utf-8', errors='ignore'))
            metadata = {}
            state["conversion_method"] = "text-inference"
            
        else:
            markdown = content.decode('utf-8', errors='ignore')
            metadata = {}
            state["conversion_method"] = "fallback-text"
            state["warnings"] = [f"ì•Œ ìˆ˜ ì—†ëŠ” íŒŒì¼ íƒ€ì…: {file_type}, í…ìŠ¤íŠ¸ë¡œ ì²˜ë¦¬"]
        
        # ğŸ”¥ ë¬¸ì„œ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ (ì˜µì…˜)
        if state.get("use_llm_metadata", False):
            doc_meta = extract_document_metadata(markdown, filename)
            metadata.update(doc_meta)
        else:
            print(f"   â„¹ï¸ LLM ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ë¹„í™œì„±í™” (ê¸°ë³¸ ë©”íƒ€ë°ì´í„°ë§Œ ì‚¬ìš©)")
        
        state["markdown"] = markdown
        state["metadata"].update(metadata)
        
    except Exception as e:
        state["errors"] = [f"ë³€í™˜ ì‹¤íŒ¨: {str(e)}"]
        state["markdown"] = ""
    
    return state


def node_convert_fallback(state: PipelineState) -> PipelineState:
    """
    2-1ë‹¨ê³„: ë³€í™˜ ì‹¤íŒ¨ ì‹œ í´ë°± ì „ëµ
    """
    content = state["content"]
    file_type = state["file_type"]
    
    state["warnings"] = [f"ê¸°ë³¸ ë³€í™˜ ì‹¤íŒ¨, í´ë°± ì „ëµ ì‹œë„ ì¤‘..."]
    
    try:
        if file_type == 'pdf':
            markdown = _pdf_fallback_extract(content)
            state["conversion_method"] = "pdf-fallback"
            
        elif file_type == 'docx':
            markdown = _docx_fallback_extract(content)
            state["conversion_method"] = "docx-fallback"
            
        else:
            markdown = content.decode('utf-8', errors='ignore')
            state["conversion_method"] = "binary-text"
        
        state["markdown"] = markdown
        state["errors"] = []
        
    except Exception as e:
        state["errors"] = [f"í´ë°± ë³€í™˜ë„ ì‹¤íŒ¨: {str(e)}"]
        state["success"] = False
    
    return state


def node_validate(state: PipelineState) -> PipelineState:
    """
    3ë‹¨ê³„: ë§ˆí¬ë‹¤ìš´ í’ˆì§ˆ ê²€ì¦
    """
    markdown = state.get("markdown", "")
    
    # ğŸ”¥ [ìµœì í™”] ë³€í™˜ì´ ì™„ë£Œë˜ì—ˆìœ¼ë¯€ë¡œ ëŒ€ìš©ëŸ‰ ë°”ì´ë„ˆë¦¬ ë°ì´í„°ëŠ” ì‚­ì œ (UI ì†ë„ í–¥ìƒ)
    if markdown and "content" in state:
        state["content"] = b"" # ë©”ëª¨ë¦¬ ë° UI ë Œë”ë§ ë¶€í•˜ ê°ì†Œ
    
    if not markdown:
        state["quality_score"] = 0.0
        state["errors"] = ["ë§ˆí¬ë‹¤ìš´ ë³€í™˜ ê²°ê³¼ê°€ ë¹„ì–´ìˆìŠµë‹ˆë‹¤."]
        return state
    
    score = 0.0
    issues = []
    
    # 1. ê¸¸ì´ ì²´í¬ (ìµœì†Œ 100ì)
    if len(markdown) >= 100:
        score += 0.2
    else:
        issues.append("í…ìŠ¤íŠ¸ê°€ ë„ˆë¬´ ì§§ìŒ")
    
    # 2. í—¤ë” ì¡´ì¬ ì—¬ë¶€
    header_count = len(re.findall(r'^#{1,6}\s+', markdown, re.MULTILINE))
    if header_count >= 3:
        score += 0.3
    elif header_count >= 1:
        score += 0.15
        issues.append("í—¤ë”ê°€ ë¶€ì¡±í•¨")
    else:
        issues.append("í—¤ë”ê°€ ì—†ìŒ")
    
    # 3. ë¬¸ë‹¨ êµ¬ì¡°
    paragraphs = [p for p in markdown.split('\n\n') if p.strip()]
    if len(paragraphs) >= 5:
        score += 0.2
    elif len(paragraphs) >= 2:
        score += 0.1
        issues.append("ë¬¸ë‹¨ êµ¬ì¡°ê°€ ë¶€ì‹¤í•¨")
    
    # 4. í•œê¸€ ë¹„ìœ¨
    korean_chars = len(re.findall(r'[ê°€-í£]', markdown))
    total_chars = len(markdown)
    korean_ratio = korean_chars / total_chars if total_chars > 0 else 0
    
    if korean_ratio >= 0.1:
        score += 0.2
    else:
        issues.append("í•œê¸€ ë¹„ìœ¨ì´ ë‚®ìŒ")
    
    # 5. íŠ¹ìˆ˜ë¬¸ì ì˜¤ì—¼ ì²´í¬
    garbage_ratio = len(re.findall(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', markdown)) / len(markdown) if markdown else 0
    if garbage_ratio < 0.01:
        score += 0.1
    else:
        issues.append("íŠ¹ìˆ˜ë¬¸ì ì˜¤ì—¼ ê°ì§€")
    
    state["quality_score"] = min(score, 1.0)
    
    if issues:
        state["warnings"] = issues
    
    return state


def node_repair(state: PipelineState) -> PipelineState:
    """
    3-1ë‹¨ê³„: ë§ˆí¬ë‹¤ìš´ í’ˆì§ˆ ë³´ì •
    """
    markdown = state.get("markdown", "")
    
    state["warnings"] = ["í’ˆì§ˆ ë³´ì • ìˆ˜í–‰ ì¤‘..."]
    state["retry_count"] = state.get("retry_count", 0) + 1
    
    # 1. íŠ¹ìˆ˜ë¬¸ì ì œê±°
    markdown = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', markdown)
    
    # 2. ì—°ì† ë¹ˆ ì¤„ ì •ë¦¬
    markdown = re.sub(r'\n{3,}', '\n\n', markdown)
    
    # 3. í—¤ë”ê°€ ì—†ìœ¼ë©´ ì¶”ë¡ í•´ì„œ ì¶”ê°€
    if not re.search(r'^#{1,6}\s+', markdown, re.MULTILINE):
        markdown = _infer_headers(markdown)
    
    # 4. ê¹¨ì§„ í…Œì´ë¸” ë³µêµ¬ ì‹œë„
    markdown = _repair_tables(markdown)
    
    state["markdown"] = markdown
    state["conversion_method"] += "+repaired"
    
    # í’ˆì§ˆ ì¬ì¸¡ì •
    state = node_validate(state)
    
    return state


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ì¡°í•­ ì¶”ì¶œ ì—”ì§„ (Pattern Registry)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

CLAUSE_PATTERNS = [
    # 1. ìˆ«ì ê³„ì¸µí˜• (1., 1.1, 5.2.1)
    {
        "name": "numeric_dot",
        "pattern": re.compile(r'^(\d+(?:\.\d+)*)\.?\s+(.*?)$'), # ë§ˆì¹¨í‘œ ë’¤ ê³µë°± í•„ìˆ˜ (1. ì œëª©)
        "level_func": lambda m: m.count('.')
    },
    # 1.1 ìˆ«ì ë‹¨ìˆœí˜• (ë§ˆì¹¨í‘œ ì—†ëŠ” ë²ˆí˜¸ + ì œëª©) -> 1 ì œëª©
    {
        "name": "numeric_simple",
        "pattern": re.compile(r'^(\d+)\s+([ê°€-í£\w].*)$'), 
        "level_func": lambda m: 0
    },
    # 2. í•œêµ­ì–´ ë²•ë ¹/ê·œì •í˜• (ì œ1ì¡°, ì œ12ì¡°)
    {
        "name": "ko_article",
        "pattern": re.compile(r'^(ì œ\s*\d+\s*ì¡°)\s*(.*?)$'),
        "level_func": lambda m: 0
    },
    # 3. í•œêµ­ì–´ ìˆœì„œí˜• (ê°€., ë‚˜., ë‹¤.)
    {
        "name": "ko_alphabet",
        "pattern": re.compile(r'^([ê°€-í£])\.\s*(.*?)$'),
        "level_func": lambda m: 1
    },
    # 4. ê´„í˜¸ ìˆ«ìí˜• ((1), (2) ë˜ëŠ” 1), 2))
    {
        "name": "bracket_numeric",
        "pattern": re.compile(r'^\(?(\d+)\)\s*(.*?)$'),
        "level_func": lambda m: 2
    },
    # 5. ì›ë¬¸ì ìˆ«ìí˜• (â‘ , â‘¡)
    {
        "name": "circle_numeric",
        "pattern": re.compile(r'^([â‘ -â‘³])\s*(.*?)$'),
        "level_func": lambda m: 2
    },
    # 6. ì˜ë¬¸ ëŒ€ë¬¸ìí˜• (A., B., C.)
    {
        "name": "en_uppercase",
        "pattern": re.compile(r'^([A-Z])\.\s*(.*?)$'),
        "level_func": lambda m: 1
    }
]

def split_by_clause(markdown: str, max_level: int = 0) -> List[Dict]:
    """
    ìœ ì—°í•œ ì¡°í•­ ì¶”ì¶œ ì—”ì§„: ì •ê·œì‹ ë ˆì§€ìŠ¤íŠ¸ë¦¬ë¥¼ ìˆœíšŒí•˜ë©° ë§¤ì¹­ë˜ëŠ” ëª¨ë“  ì¡°í•­ ë¶„í• 
    """
    lines = markdown.split('\n')
    clauses = []
    
    # í˜„ì¬ ë§¤ì¹­ëœ íŒ¨í„´ ì´ë¦„ (ë¬¸ì„œ ë‚´ ì¼ê´€ì„± ìœ ì§€ ì‹œë„)
    active_pattern_names = set()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # í—¤ë” ë§ˆì»¤ ì œê±°
        clean_line = re.sub(r'^#+\s*', '', line)
        
        # íŒ¨í„´ ë§¤ì¹­ ì‹œë„
        match = None
        matched_pattern = None
        for p in CLAUSE_PATTERNS:
            m = p["pattern"].match(clean_line)
            if m:
                # ì¡°í•­ ë²ˆí˜¸ ë’¤ì— ì œëª©ì´ ì—†ë”ë¼ë„ ì¸ì‹ (ë²ˆí˜¸ë§Œ ìˆëŠ” ë¼ì¸ ë°©ì–´)
                num = m.group(1)
                title = m.group(2).strip() or f"Section {num}"
                match = m
                matched_pattern = p
                break
        
        if match:
            clause_num = match.group(1)
            title = match.group(2).strip() or f"Section {clause_num}"
            level = matched_pattern["level_func"](clause_num)
            
            # max_level í•„í„° (0ì´ë©´ ë¬´ì œí•œ)
            if max_level > 0 and level > max_level:
                i += 1
                continue
                
            content_lines = []
            i += 1
            
            # ë‹¤ìŒ ì¡°í•­ì´ ë‚˜ì˜¬ ë•Œê¹Œì§€ ìˆ˜ì§‘
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    content_lines.append(lines[i])
                    i += 1
                    continue
                    
                next_clean = re.sub(r'^#+\s*', '', next_line)
                
                # ë‹¤ìŒ ë¼ì¸ì´ ì–´ë–¤ ì¡°í•­ íŒ¨í„´ì— ê±¸ë¦¬ëŠ”ì§€ í™•ì¸
                is_next_clause = False
                for p in CLAUSE_PATTERNS:
                    if p["pattern"].match(next_clean):
                        is_next_clause = True
                        break
                
                if is_next_clause:
                    break
                    
                content_lines.append(lines[i])
                i += 1
            
            content = '\n'.join(content_lines).strip()
            
            # ê³„ì¸µ êµ¬ì¡° (ìˆ«ìí˜•ì¼ ë•Œë§Œ ë¶€ëª¨ ê³„ì‚°)
            parent_clauses = []
            if matched_pattern["name"] == "numeric_dot":
                parts = clause_num.split('.')
                for j in range(len(parts) - 1):
                    parent_clauses.append('.'.join(parts[:j+1]))
            
            clauses.append({
                "clause": clause_num,
                "title": title,
                "content": content,
                "level": level,
                "parent_clauses": parent_clauses,
                "pattern": matched_pattern["name"]
            })
        else:
            i += 1
            
    return clauses


def detect_content_start_with_llm(markdown: str) -> Optional[str]:
    """
    LLMì„ ì‚¬ìš©í•˜ì—¬ ë¬¸ì„œ í—¤ë”/ëª©ì°¨ êµ¬ê°„ì´ ëë‚˜ê³  ì‹¤ì œ ë³¸ë¬¸ ì¡°í•­ì´ ì‹œì‘ë˜ëŠ” ì§€ì ì˜ í…ìŠ¤íŠ¸(Anchor)ë¥¼ ì°¾ìŠµë‹ˆë‹¤.
    (ì˜ˆ: "5.10.3", "ì œ1ì¡°" ë“±)
    """
    sample_text = markdown[:5000]
    
    prompt = f"""ë‹¹ì‹ ì€ GMP ë¬¸ì„œ êµ¬ì¡° ë¶„ì„ ì „ë¬¸ê°€ì…ë‹ˆë‹¤.
ë‹¤ìŒ ë§ˆí¬ë‹¤ìš´ ë¬¸ì„œ ë‚´ìš©ì„ ë³´ê³ , ë¬¸ì„œ ì •ë³´(SOP Number, Version ë“±)ì™€ ëª©ì°¨(ToC)ê°€ ì™„ì „íˆ ì¢…ë£Œë˜ê³ ,
ì‹¤ì œ ì‹¤í–‰ ì§€ì¹¨ì´ë‚˜ ì¡°í•­(Clause)ì´ ë³¸ê²©ì ìœ¼ë¡œ ì‹œì‘ë˜ëŠ” ì²« ë²ˆì§¸ ë¬¸ì¥ì´ë‚˜ ì¡°í•­ ë²ˆí˜¸ë¥¼ ì°¾ìœ¼ì„¸ìš”.

[ì‘ì—… ê·œì¹™]
1. ë³¸ë¬¸ì´ ì‹œì‘ë˜ëŠ” ì§€ì ì˜ í…ìŠ¤íŠ¸(ì˜ˆ: "5.1", "ì œ1ì¡°", "1. ëª©ì ")ë¥¼ ì •í™•íˆ í•œ ë¬¸ì¥ë§Œ ì¶œë ¥í•˜ì„¸ìš”.
2. ë§Œì•½ ë¬¸ì„œì˜ ì²˜ìŒë¶€í„° ë°”ë¡œ ë³¸ë¬¸ì´ ì‹œì‘ëœë‹¤ë©´ "START"ë¼ê³  ë‹µë³€í•˜ì„¸ìš”.
3. ìƒê° ê³¼ì • ì—†ì´ ê²°ê³¼ í…ìŠ¤íŠ¸ë§Œ ì¶œë ¥í•˜ì„¸ìš”.

[ë¬¸ì„œ ìƒ˜í”Œ]
{sample_text}
"""
    try:
        # Z.AI: í•œêµ­ì–´ ë¬¸ì„œ êµ¬ì¡° ë° ì¡°í•­ ì‹œì‘ ì§€ì  íŒŒì•…ì— íƒì›”
        anchor = get_llm_response(prompt, llm_backend="zai", max_tokens=100, temperature=0.1).strip()
        if not anchor or "START" in anchor.upper():
            return None
        # Anchorê°€ ë„ˆë¬´ ê¸¸ë©´ ë¬´ì‹œ (ì •í™•í•œ ë§¤ì¹­ì„ ìœ„í•´)
        if len(anchor) > 100:
            return None
        return anchor
    except:
        return None

def discover_structure_with_llm(markdown: str) -> List[Dict]:
    """
    ì •ê·œì‹ ë§¤ì¹­ì´ ì‹¤íŒ¨í•  ê²½ìš°, LLMì„ ì‚¬ìš©í•˜ì—¬ ë¬¸ì„œ ì „ì²´ì˜ ì¡°í•­ë“¤ì„ ê³„ì¸µì ìœ¼ë¡œ ì°¾ì•„ëƒ…ë‹ˆë‹¤.
    """
    # ë¬¸ì„œê°€ ê¸¸ ê²½ìš° ë¶„í•  ë¶„ì„ì´ í•„ìš”í•  ìˆ˜ ìˆìœ¼ë‚˜, ì¼ë‹¨ ë„‰ë„‰í•˜ê²Œ ë¶„ì„
    sample_text = markdown[:15000] 
    
    prompt = f"""ë‹¹ì‹ ì€ ì „ë¬¸ ë¬¸ì„œ êµ¬ì¡°í™” ì—”ì§„ì…ë‹ˆë‹¤.
ë‹¤ìŒ ë§ˆí¬ë‹¤ìš´ ë¬¸ì„œì—ì„œ 'ëª¨ë“  ì¡°í•­(Clause)' ë˜ëŠ” 'ì„¹ì…˜'ì„ ëˆ„ë½ ì—†ì´ ì°¾ì•„ë‚´ì–´ JSON ë¦¬ìŠ¤íŠ¸ë¡œ ì‘ë‹µí•˜ì„¸ìš”.

[ì¶œë ¥ í˜•ì‹]
[
  {{"clause": "ë²ˆí˜¸", "title": "ì œëª©(í•œê¸€/ì˜ë¬¸ í¬í•¨)", "content": "í•´ë‹¹ ì¡°í•­ì˜ ì „ë¬¸"}},
  ...
]

[ì‘ì—… ê·œì¹™]
1. ë¬¸ì„œì— ì¡´ì¬í•˜ëŠ” ëª¨ë“  ì¡°í•­ì„ ë¹ ì§ì—†ì´ í¬í•¨í•˜ì„¸ìš”. (ê°œìˆ˜ ì œí•œ ì—†ìŒ)
2. ë²ˆí˜¸ê°€ ì—†ëŠ” ë‹¨ë½ì€ ì œëª©ì˜ í•µì‹¬ í‚¤ì›Œë“œë¥¼ ë²ˆí˜¸ ëŒ€ì‹  ì‚¬ìš©í•˜ê±°ë‚˜ ìƒëµí•˜ì„¸ìš”.
3. í•œê¸€ê³¼ ì˜ë¬¸ì´ í˜¼ìš©ëœ ê²½ìš° ì œëª©ì— ë‘ ì–¸ì–´ë¥¼ ëª¨ë‘ í¬í•¨í•˜ì„¸ìš”.
4. JSON ë°ì´í„° ì™¸ì— ì–´ë–¤ í…ìŠ¤íŠ¸ë„ ì¶œë ¥í•˜ì§€ ë§ˆì„¸ìš”.

[ë¬¸ì„œ ë‚´ìš©]
{sample_text}
"""
    try:
        # Z.AI: ë³µì¡í•œ SOP ê³„ì¸µ êµ¬ì¡°ë¥¼ ëˆ„ë½ ì—†ì´ ë¶„ì„í•˜ëŠ” ë° ìœ ë¦¬
        llm_res = get_llm_response(prompt, llm_backend="zai", max_tokens=4000, temperature=0.1)
        json_match = re.search(r'\[.*\]', llm_res, re.DOTALL)
        if json_match:
            discovered = json.loads(json_match.group(0))
            for item in discovered:
                item["level"] = 0
                item["parent_clauses"] = []
                item["pattern"] = "ai_discovered"
            return discovered
    except:
        return []
    return []

def node_split(state: PipelineState) -> PipelineState:
    """
    4ë‹¨ê³„: ì¡°í•­ ë˜ëŠ” í—¤ë” ê¸°ì¤€ ë¶„í•  + ê³„ì¸µ êµ¬ì¡° êµ¬ì¶•
    ğŸ”¥ v9.5: SOP ìµœì í™” (ê°œë³„ ì¡°í•­ ì •ë°€ íŒŒì‹±)
    """
    markdown = state.get("markdown", "")
    use_clause_parsing = state.get("use_clause_parsing", True)

    if not markdown:
        state["sections"] = []
        return state

    # 1. ì§€ëŠ¥í˜• ì‹œì‘ ì§€ì  ì‹ë³„ (í—¤ë”/ëª©ì°¨ ìŠ¤í‚µ)
    start_anchor = detect_content_start_with_llm(markdown)
    effective_markdown = markdown
    if start_anchor:
        anchor_idx = markdown.find(start_anchor)
        if anchor_idx >= 0:
            print(f"   ğŸ¯ [ToC/Header Skip] '{start_anchor[:20]}...' ì§€ì ë¶€í„° ë³¸ë¬¸ íŒŒì‹±ì„ ì‹œì‘í•©ë‹ˆë‹¤.")
            effective_markdown = markdown[anchor_idx:]

    # 2. ì¡°í•­ ë²ˆí˜¸ ê¸°ë°˜ ì •ê·œì‹ íŒŒì‹± ì‹œë„
    if use_clause_parsing:
        sections = split_by_clause(effective_markdown, max_level=0)
        
        # ğŸ”¥ [Structure Discovery Fallback] ì •ê·œì‹ìœ¼ë¡œ í•˜ë‚˜ë„ ì•ˆ ì¡í ë•Œë§Œ AI ë™ì›
        if not sections and len(effective_markdown) > 300:
            print("   ğŸ” [Structure Discovery] ì •ê·œì‹ ë§¤ì¹­ ì‹¤íŒ¨, AI ê¸°ë°˜ ì •ë°€ êµ¬ì¡° ë¶„ì„ ìˆ˜í–‰ ì¤‘...")
            sections = discover_structure_with_llm(effective_markdown)
            
        if sections:
            # í›„ì²˜ë¦¬: ê³„ì¸µ êµ¬ì¡°(Graph DBìš©) ë° ë©”íƒ€ë°ì´í„° ë³´ê°•
            for section in sections:
                section["page"] = 1 # ê¸°ë³¸ê°’
                # ìƒìœ„ ê´€ê³„ ì¶”ì¶œ (ì •ê·œì‹ íŒŒì‹± ê²°ê³¼ì— parent_clausesê°€ ìˆìŒ)
                if not section.get("parent") and section.get("parent_clauses"):
                    section["parent"] = section["parent_clauses"][-1]
                
                # ğŸ”¥ [V22.1 Normalization] headers ê°€ ì—†ëŠ” ê²½ìš° (split_by_clause/AI ê²°ê³¼) ë³´ì¶©
                if "headers" not in section:
                    clause_num = section.get("clause", "")
                    title = section.get("title", "Untitled")
                    full_title = f"{clause_num} {title}".strip() if clause_num else title
                    level = section.get("level", 0)
                    h_level = min(6, level + 1) # Level 0 -> H1, Level 1 -> H2
                    
                    section["headers"] = {f"H{h_level}": full_title}
                    section["header_path"] = full_title
                    # ìƒìœ„ ê²½ë¡œê°€ ìˆë‹¤ë©´ ë” ë³´ê°• ê°€ëŠ¥í•˜ê² ìœ¼ë‚˜ í˜„ì¬ëŠ” ì´ì •ë„ë©´ ì¶©ë¶„
                
            state["sections"] = sections
            print(f"   âœ… [Split] {len(sections)}ê°œì˜ ê°œë³„ ì¡°í•­ì´ ì¶”ì¶œë˜ì—ˆìŠµë‹ˆë‹¤.")
            return state

    # 3. í—¤ë” ê¸°ë°˜ íŒŒì‹± (í´ë°±)
    lines = effective_markdown.split('\n')
    sections = []
    current_headers = {1: None, 2: None, 3: None, 4: None, 5: None, 6: None}
    current_content = []
    current_page = 1
    
    def flush_section():
        nonlocal current_content
        if current_content:
            content = '\n'.join(current_content).strip()
            if content:
                header_path_parts = []
                headers_dict = {}
                for level in range(1, 7):
                    if current_headers[level]:
                        headers_dict[f"H{level}"] = current_headers[level]
                        header_path_parts.append(current_headers[level])
                
                parent = None
                for level in range(6, 0, -1):
                    if current_headers[level]:
                        for p_level in range(level - 1, 0, -1):
                            if current_headers[p_level]:
                                parent = current_headers[p_level]
                                break
                        break
                
                sections.append({
                    "content": content,
                    "headers": headers_dict,
                    "header_path": " > ".join(header_path_parts) if header_path_parts else None,
                    "page": current_page,
                    "parent": parent,
                })
        current_content = []
    
    for line in lines:
        page_match = re.match(r'<!-- PAGE:(\d+) -->', line)
        if page_match:
            current_page = int(page_match.group(1))
            continue
        
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            flush_section()
            level = len(header_match.group(1))
            header_text = header_match.group(2).strip()
            current_headers[level] = header_text
            for l in range(level + 1, 7):
                current_headers[l] = None
            current_content.append(line)
        else:
            current_content.append(line)
    
    flush_section()
    state["sections"] = sections
    return state


def node_optimize(state: PipelineState) -> PipelineState:
    """
    5ë‹¨ê³„: ê¸´ ì„¹ì…˜ ì¬ë¶„í•  + ìƒì„¸ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ + ìµœì í™”
    """
    sections = state.get("sections", [])
    chunk_size = state.get("chunk_size", 500)
    chunk_overlap = state.get("chunk_overlap", 50)
    doc_meta = state.get("metadata", {})
    use_llm_metadata = state.get("use_llm_metadata", False)
    
    chunks = []
    idx = 0
    
    # ğŸ”¥ ë¬¸ì„œ ë ˆë²¨ ë©”íƒ€ë°ì´í„°
    doc_id = doc_meta.get("doc_id")
    doc_title = doc_meta.get("title")
    version = doc_meta.get("version")
    
    # 1ë‹¨ê³„: ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ë° ì¡°í•­ ë¶„ì„ (ì•ˆì •ì„±ì„ ìœ„í•´ ë‹¤ì‹œ ìˆœì°¨ ì²˜ë¦¬)
    for section in sections:
        content = section["content"]
        headers = section.get("headers", {})
        clause_level = 0
        for l in range(6, 0, -1):
            if headers.get(f"H{l}"):
                clause_level = l
                break
        
        # ì¡°í•­ ë²ˆí˜¸ ë° ì œëª© ì¶”ì¶œ
        current_section_title = headers.get(f"H{clause_level}") or "Untitled"
        clause_id = None
        num_match = re.search(r'([ì œ]?\d+(?:\.\d+)*[ì¡°]?)', current_section_title)
        if num_match:
            clause_id = num_match.group(1)
            
        # ìƒìœ„ ì„¹ì…˜ ë²ˆí˜¸ ìœ ì¶” (5.1.2 -> 5)
        main_section = clause_id.split('.')[0] if clause_id and '.' in clause_id else clause_id
        
        # ì„¹ì…˜ ì •ë³´ ì €ì¥
        section["main_section"] = main_section
        section["clause_id"] = clause_id
        section["current_title"] = current_section_title
        section["clause_level"] = clause_level

        # ğŸ”¥ ì¡°í•­ë³„ ìƒì„¸ ë©”íƒ€ë°ì´í„° ì¶”ì¶œ (AI) - ìˆœì°¨ ì²˜ë¦¬ë¡œ ë³µêµ¬
        # ì œëª©ì´ ìˆê³ , ì¡°í•­ ë²ˆí˜¸ê°€ ìˆìœ¼ë©°, ë³¸ë¬¸ì´ 100ì ì´ìƒì¸ ê²½ìš°ì—ë§Œ ë¶„ì„
        clause_meta = {}
        if use_llm_metadata and current_section_title != "Untitled" and clause_id and len(content.strip()) > 100:
            print(f"   ğŸ” [{sections.index(section)+1}/{len(sections)}] ì¡°í•­ ë¶„ì„ ì¤‘: {current_section_title}")
            clause_meta = extract_clause_metadata(content, doc_meta, current_section_title)

        section["clause_meta"] = clause_meta
        idx += 1 # ğŸ”¥ ë‹¤ìŒ ì„¹ì…˜ì„ ìœ„í•´ ì¸ë±ìŠ¤ ì¦ê°€ (ëˆ„ë½ ë°©ì§€)

    # 2ë‹¨ê³„: ìµœì í™” ë° ì²­í¬ ìƒì„±
    for section in sections:
        content = section["content"]
        clause_id = section.get("clause_id")
        current_section_title = section.get("current_title")
        clause_meta = section.get("clause_meta", {})
        main_section = section.get("main_section")
        clause_level = section.get("clause_level", 0)
        
        # ê¸´ ì„¹ì…˜ ì¬ë¶„í• 
        if len(content) > chunk_size:
            text_chunks = _split_recursive(content, chunk_size, chunk_overlap)
            is_split = len(text_chunks) > 1
        else:
            text_chunks = [content]
            is_split = False
        
        for i, text in enumerate(text_chunks):
            if not text.strip(): continue
            
            section_id = f"{doc_id}:{clause_id}" if clause_id else f"{doc_id}:CH{idx}"
            
            # ğŸ”¥ ê³ ë„í™”ëœ ë©”íƒ€ë°ì´í„° êµ¬ì¡° (V22.0)
            meta = {
                "doc_id": doc_id,
                "doc_title": doc_title,
                "clause_id": clause_id,
                "title": current_section_title,
                "clause_level": clause_level,
                "main_section": main_section,
                "section_id": section_id,
                # LLM ë¶„ì„ ë°ì´í„°
                "content_type": clause_meta.get("content_type"),
                "main_topic": clause_meta.get("main_topic"),
                "sub_topics": clause_meta.get("sub_topics"),
                "actors": clause_meta.get("actors"),
                "actions": clause_meta.get("actions"),
                "conditions": clause_meta.get("conditions"),
                "summary": clause_meta.get("summary"),
                "intent_scope": clause_meta.get("intent_scope"),
                "intent_summary": clause_meta.get("intent_summary"),
                "language": clause_meta.get("language", "ko"),
                # ì‹œìŠ¤í…œ ê´€ë¦¬ ì •ë³´
                "page": section.get("page", 1),
                "parent": section.get("parent"),
                "chunk_part": i + 1 if is_split else None,
            }
            
            chunks.append({
                "text": text.strip(),
                "index": idx,
                "metadata": meta
            })
            idx += 1
    
    state["chunks"] = chunks
    return state


def node_finalize(state: PipelineState) -> PipelineState:
    """
    6ë‹¨ê³„: ìµœì¢… ê²°ê³¼ ì •ë¦¬
    """
    chunks = state.get("chunks", [])
    
    if not chunks:
        state["success"] = False
        state["errors"] = ["ì²­í¬ ìƒì„± ì‹¤íŒ¨: ê²°ê³¼ê°€ ë¹„ì–´ìˆìŠµë‹ˆë‹¤."]
    else:
        state["success"] = True
    
    # í†µê³„ ì¶”ê°€
    state["metadata"]["total_chunks"] = len(chunks)
    state["metadata"]["quality_score"] = state.get("quality_score", 0)
    state["metadata"]["conversion_method"] = state.get("conversion_method", "unknown")
    
    return state


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ì¡°ê±´ë¶€ ë¼ìš°íŒ… í•¨ìˆ˜
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def should_fallback(state: PipelineState) -> Literal["fallback", "validate"]:
    """ë³€í™˜ ì‹¤íŒ¨ ì‹œ í´ë°±ìœ¼ë¡œ ë¼ìš°íŒ…"""
    if state.get("errors") or not state.get("markdown"):
        return "fallback"
    return "validate"


def should_repair(state: PipelineState) -> Literal["repair", "split"]:
    """í’ˆì§ˆ ì ìˆ˜ê°€ ë‚®ìœ¼ë©´ ë³´ì •ìœ¼ë¡œ ë¼ìš°íŒ…"""
    quality_score = state.get("quality_score", 0)
    retry_count = state.get("retry_count", 0)
    
    if quality_score < 0.5 and retry_count < 2:
        return "repair"
    return "split"


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# í—¬í¼ í•¨ìˆ˜ë“¤
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def _convert_docx(filename: str, content: bytes) -> tuple:
    """DOCX â†’ Markdown"""
    from docx import Document
    
    doc = Document(BytesIO(content))
    md_lines = []
    metadata = {}
    
    sop_pattern = re.compile(r'((?:EQ-)?SOP[-_]?\d{4,5})', re.IGNORECASE)
    
    main_sections = ['ëª©ì ', 'Purpose', 'ì ìš© ë²”ìœ„', 'Scope', 'ì •ì˜', 'Definitions',
                     'ì±…ì„', 'Responsibilities', 'ì ˆì°¨', 'Procedure', 
                     'ì°¸ê³ ë¬¸í—Œ', 'Reference', 'ì²¨ë¶€', 'Attachments']
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
        
        # SOP ID ì¶”ì¶œ
        sop_match = sop_pattern.search(text)
        if sop_match and "sop_id" not in metadata:
            sop_id = sop_match.group(1).upper().replace('_', '-')
            if not sop_id.startswith('EQ-'):
                sop_id = 'EQ-' + sop_id
            metadata["sop_id"] = sop_id
        
        # í—¤ë” ë ˆë²¨ ê²°ì •
        header_level = None
        
        style_name = para.style.name.lower() if para.style else ""
        if 'heading 1' in style_name or 'title' in style_name:
            header_level = 1
        elif 'heading 2' in style_name:
            header_level = 2
        elif 'heading 3' in style_name:
            header_level = 3
        elif 'heading 4' in style_name:
            header_level = 4
        
        if not header_level:
            for section in main_sections:
                if text.startswith(section) or re.match(rf'^\d+\s+{section}', text):
                    header_level = 2
                    break
            
            if not header_level:
                if re.match(r'^\d+\.\d+\.\d+\.\d+\s*', text) and len(text) < 60:
                    header_level = 5
                elif re.match(r'^\d+\.\d+\.\d+\s+', text) and len(text) < 60:
                    header_level = 4
                elif re.match(r'^\d+\.\d+\s+', text) and len(text) < 60:
                    header_level = 3
                elif re.match(r'^\d+\.?\s+[ê°€-í£A-Za-z]', text) and len(text) < 60:
                    header_level = 2
                elif re.match(r'^[ê°€-í£A-Z][ê°€-í£\s\(\)/Â·\-]+\s*\([A-Za-z\s&/\-:]+\)\s*$', text):
                    header_level = 3
        
        if header_level:
            md_lines.append(f"{'#' * header_level} {text}")
        else:
            md_lines.append(text)
    
    # í…Œì´ë¸” ì²˜ë¦¬
    for table in doc.tables:
        md_lines.append("")
        md_lines.append(_table_to_markdown(table))
    
    return '\n'.join(md_lines), metadata


def _convert_pdf_with_fallback(filename: str, content: bytes) -> tuple:
    """PDF ë³€í™˜ (ë‹¤ì¤‘ í´ë°±) + í˜ì´ì§€ ë§ˆì»¤"""
    
    # 1ìˆœìœ„: pdfplumber (ê°€ì¥ ì•ˆì •ì )
    try:
        import pdfplumber
        md_lines = []
        total_text_len = 0
        with pdfplumber.open(BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                if text.strip():
                    md_lines.append(f"<!-- PAGE:{i + 1} -->")
                    md_lines.append(text)
                    total_text_len += len(text.strip())
        
        # ğŸ”¥ ì‹¤ì œ í…ìŠ¤íŠ¸ê°€ ì˜ë¯¸ ìˆëŠ” ìˆ˜ì¤€(ì˜ˆ: 50ì) ì´ìƒì¼ ë•Œë§Œ ì„±ê³µìœ¼ë¡œ ê°„ì£¼
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pdfplumber", "total_pages": len(pdf.pages)}, "pdfplumber"
        print(f"   âš ï¸ pdfplumber: í…ìŠ¤íŠ¸ ì¶”ì¶œ ë¶€ì¡± ({total_text_len}ì). ë‹¤ë¥¸ íŒŒì„œ ì‹œë„...")
    except Exception as e:
        print(f"   pdfplumber ì‹¤íŒ¨: {e}")
    
    # 2ìˆœìœ„: Docling
    try:
        from docling.document_converter import DocumentConverter
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(content)
            temp_path = f.name
        
        try:
            converter = DocumentConverter()
            result = converter.convert(temp_path)
            markdown = result.document.export_to_markdown()
            if len(markdown.strip()) > 50:
                return markdown, {"parser": "docling"}, "docling"
            print(f"   âš ï¸ Docling: í…ìŠ¤íŠ¸ ì¶”ì¶œ ë¶€ì¡±. ë‹¤ë¥¸ íŒŒì„œ ì‹œë„...")
        finally:
            os.unlink(temp_path)
    except Exception as e:
        print(f"   Docling ì‹¤íŒ¨: {e}")
    
    # 3ìˆœìœ„: PyMuPDF
    try:
        import fitz
        pdf = fitz.open(stream=content, filetype="pdf")
        md_lines = []
        total_text_len = 0
        for page_num, page in enumerate(pdf):
            text = page.get_text()
            if text.strip():
                md_lines.append(f"<!-- PAGE:{page_num + 1} -->")
                md_lines.append(text)
                total_text_len += len(text.strip())
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pymupdf", "total_pages": len(pdf)}, "pymupdf"
        print(f"   âš ï¸ PyMuPDF: í…ìŠ¤íŠ¸ ì¶”ì¶œ ë¶€ì¡± ({total_text_len}ì). ë‹¤ë¥¸ íŒŒì„œ ì‹œë„...")
    except Exception as e:
        print(f"   PyMuPDF ì‹¤íŒ¨: {e}")
    
    # 4ìˆœìœ„: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(content))
        md_lines = []
        total_text_len = 0
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                md_lines.append(f"<!-- PAGE:{i + 1} -->")
                md_lines.append(text)
                total_text_len += len(text.strip())
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pypdf2", "total_pages": len(reader.pages)}, "pypdf2"
        print(f"   âš ï¸ PyPDF2: í…ìŠ¤íŠ¸ ì¶”ì¶œ ë¶€ì¡± ({total_text_len}ì). ë‹¤ë¥¸ íŒŒì„œ ì‹œë„...")
    except Exception as e:
        print(f"   PyPDF2 ì‹¤íŒ¨: {e}")
    
    # 5ìˆœìœ„: ğŸ”¥ OCR Fallback (ìŠ¤ìº”ë³¸/ì´ë¯¸ì§€ PDFìš©)
    try:
        print("   ğŸ” ìŠ¤ìº” ë¬¸ì„œ ê°ì§€: OCR(ê´‘í•™ ë¬¸ì ì¸ì‹) ì—”ì§„ ê°€ë™ ì¤‘...")
        markdown, metadata = _convert_pdf_ocr(content)
        if len(markdown.strip()) > 50:
            return markdown, {**metadata, "parser": "ocr"}, "ocr"
    except Exception as e:
        print(f"   OCR íŒŒì„œ ì‹¤íŒ¨: {e}")
    
    raise Exception("ëª¨ë“  PDF íŒŒì„œ ì‹¤íŒ¨ (OCR í¬í•¨)")

def _convert_pdf_ocr(content: bytes) -> tuple:
    """PDF OCR ì²˜ë¦¬ (Tesseract ê¸°ë°˜)"""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        from PIL import Image
        
        # PDFë¥¼ ì´ë¯¸ì§€ë¡œ ë³€í™˜ (300 DPI ê¶Œì¥)
        pages = convert_from_bytes(content, dpi=300)
        
        md_lines = []
        for i, page in enumerate(pages):
            # í•œê¸€ + ì˜ì–´ OCR ìˆ˜í–‰
            text = pytesseract.image_to_string(page, lang='kor+eng')
            if text.strip():
                md_lines.append(f"<!-- PAGE:{i + 1} (OCR) -->")
                md_lines.append(text)
        
        return '\n'.join(md_lines), {"total_pages": len(pages)}
    except ImportError:
        return "OCR í•„ìš”í•œ ë¼ì´ë¸ŒëŸ¬ë¦¬(pytesseract, pdf2image)ê°€ ì„¤ì¹˜ë˜ì§€ ì•Šì•˜ìŠµë‹ˆë‹¤.", {}
    except Exception as e:
        return f"OCR ì²˜ë¦¬ ì¤‘ ì˜¤ë¥˜ ë°œìƒ: {str(e)}", {}


def _convert_html(filename: str, content: bytes) -> tuple:
    """HTML â†’ Markdown"""
    from bs4 import BeautifulSoup
    
    html = content.decode('utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'html.parser')
    
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    
    md_lines = []
    title = soup.title.string if soup.title else filename
    md_lines.append(f"# {title}")
    md_lines.append("")
    
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
        if tag.name.startswith('h'):
            level = int(tag.name[1])
            md_lines.append(f"{'#' * level} {tag.get_text(strip=True)}")
        elif tag.name == 'li':
            md_lines.append(f"- {tag.get_text(strip=True)}")
        else:
            text = tag.get_text(strip=True)
            if text:
                md_lines.append(text)
        md_lines.append("")
    
    return '\n'.join(md_lines), {"title": title}


def _convert_text_to_markdown(text: str) -> str:
    """í…ìŠ¤íŠ¸ â†’ ë§ˆí¬ë‹¤ìš´ (í—¤ë” ì¶”ë¡ )"""
    return _infer_headers(text)


def _pdf_fallback_extract(content: bytes) -> str:
    """PDF í´ë°± ì¶”ì¶œ"""
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(content)) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                texts.append(f"<!-- PAGE:{i+1} -->")
                texts.append(page.extract_text() or '')
            return '\n\n'.join(texts)
    except:
        pass
    return content.decode('latin-1', errors='ignore')


def _docx_fallback_extract(content: bytes) -> str:
    """DOCX í´ë°±: XML ì§ì ‘ íŒŒì‹±"""
    import zipfile
    from xml.etree import ElementTree
    
    try:
        with zipfile.ZipFile(BytesIO(content)) as zf:
            xml_content = zf.read('word/document.xml')
            tree = ElementTree.fromstring(xml_content)
            
            texts = []
            for t in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    texts.append(t.text)
            
            return '\n'.join(texts)
    except:
        return ""


def _table_to_markdown(table) -> str:
    """Word í…Œì´ë¸” â†’ Markdown"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    md_lines = []
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
    
    return '\n'.join(md_lines)


def _infer_headers(markdown: str) -> str:
    """í—¤ë” ì¶”ë¡  ì‚½ì… (PDFìš© ê°•í™”)"""
    lines = markdown.split('\n')
    result = []
    
    main_sections = ['ëª©ì ', 'ì ìš© ë²”ìœ„', 'ì •ì˜', 'ì±…ì„', 'ì ˆì°¨', 'ì°¸ê³ ë¬¸í—Œ', 'ì²¨ë¶€',
                     'Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Procedure', 
                     'Reference', 'Attachments']
    
    # ğŸ”¥ ë¬´ì‹œí•  íŒ¨í„´ (í˜ì´ì§€ ë²ˆí˜¸ ë“±) - í…ìŠ¤íŠ¸ëŠ” ìœ ì§€í•˜ë˜ í—¤ë”ë¡œ ì•ˆ ë§Œë“¦
    ignore_patterns = [
        r'^\d+\s+of\s+\d+$',
        r'^Page\s+\d+',
        r'^-\s*\d+\s*-$',
        r'^Number:\s*',
        r'^<!--\s*PAGE',
    ]
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            result.append(line)
            continue
        
        # ë¬´ì‹œ íŒ¨í„´ ì²´í¬ (í—¤ë”ë¡œ ì•ˆ ë§Œë“¤ê³  í…ìŠ¤íŠ¸ ìœ ì§€)
        should_ignore = False
        for pattern in ignore_patterns:
            if re.match(pattern, stripped, re.IGNORECASE):
                should_ignore = True
                break
        
        if should_ignore:
            result.append(line)
            continue
        
        # 1. ìˆ«ìí˜• í—¤ë” íŒ¨í„´
        # 5.1.2.1 xxx â†’ H5 (ê¸€ì ìˆ˜ ì œí•œ ê°•í™”: 40ì)
        if re.match(r'^(\d+\.\d+\.\d+\.\d+)\s*(.+)', stripped) and len(stripped) < 40:
            result.append(f"##### {stripped}")
            continue
        
        # 5.1.1 xxx â†’ H4
        if re.match(r'^(\d+\.\d+\.\d+)\s+(.+)', stripped) and len(stripped) < 40:
            result.append(f"#### {stripped}")
            continue
        
        # 5.1 xxx â†’ H3 (ê¸€ì ìˆ˜ ì œí•œ ê°•í™”: 40ì)
        if re.match(r'^(\d+\.\d+)\s+(.+)', stripped) and len(stripped) < 40:
            result.append(f"### {stripped}")
            continue
        
        # 5 xxx â†’ H2
        match = re.match(r'^(\d+)\s+([ê°€-í£A-Za-z].+)', stripped)
        if match:
            num = match.group(1)
            text = match.group(2)
            if not re.match(r'^of\s+\d+', text, re.IGNORECASE) and len(stripped) < 40:
                result.append(f"## {stripped}")
                continue
        
        # 2. ì£¼ìš” ì„¹ì…˜ í‚¤ì›Œë“œ â†’ H2
        matched = False
        for section in main_sections:
            if stripped.startswith(section) and len(stripped) < 50:
                result.append(f"## {stripped}")
                matched = True
                break
        
        if not matched:
            # 3. ì†Œì œëª© íŒ¨í„´ â†’ H3
            if re.match(r'^[ê°€-í£][ê°€-í£\s\(\)/Â·\-]+\s*\([A-Za-z\s&/\-:]+\)\s*$', stripped):
                result.append(f"### {stripped}")
            else:
                result.append(line)
    
    return '\n'.join(result)


def _repair_tables(markdown: str) -> str:
    """ê¹¨ì§„ í…Œì´ë¸” ë³µêµ¬"""
    lines = markdown.split('\n')
    result = []
    in_table = False
    table_cols = 0
    
    for line in lines:
        if line.strip().startswith('|') and line.strip().endswith('|'):
            cols = line.count('|') - 1
            
            if not in_table:
                in_table = True
                table_cols = cols
                result.append(line)
            else:
                while line.count('|') - 1 < table_cols:
                    line = line.rstrip('|') + ' |'
                result.append(line)
        else:
            in_table = False
            result.append(line)
    
    return '\n'.join(result)


def _split_recursive(text: str, chunk_size: int, overlap: int) -> List[str]:
    """ì¬ê·€ì  í…ìŠ¤íŠ¸ ë¶„í• """
    if len(text) <= chunk_size:
        return [text]
    
    separators = ["\n\n", "\n| ", "\n", ". ", "ã€‚", " ", ""]
    is_table = text.strip().startswith('|') or '\n|' in text
    effective_overlap = 0 if is_table else overlap
    
    for sep in separators:
        if sep in text:
            parts = text.split(sep)
            chunks = []
            current = ""
            
            for part in parts:
                if len(current) + len(part) + len(sep) <= chunk_size:
                    current = current + sep + part if current else part
                else:
                    if current:
                        chunks.append(current)
                    if len(part) > chunk_size:
                        chunks.extend(_split_recursive(part, chunk_size, overlap))
                        current = ""
                    else:
                        current = part
            
            if current:
                chunks.append(current)
            
            return chunks
    
    step = chunk_size - effective_overlap if effective_overlap > 0 else chunk_size
    return [text[i:i+chunk_size] for i in range(0, len(text), step)]


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# LangGraph íŒŒì´í”„ë¼ì¸ ë¹Œë”
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def build_pipeline():
    """LangGraph íŒŒì´í”„ë¼ì¸ êµ¬ì„±"""
    try:
        from langgraph.graph import StateGraph, END
    except ImportError:
        raise ImportError("langgraph íŒ¨í‚¤ì§€ê°€ í•„ìš”í•©ë‹ˆë‹¤: pip install langgraph")
    
    workflow = StateGraph(PipelineState)
    
    workflow.add_node("load", node_load)
    workflow.add_node("convert", node_convert)
    workflow.add_node("fallback", node_convert_fallback)
    workflow.add_node("validate", node_validate)
    workflow.add_node("repair", node_repair)
    workflow.add_node("split", node_split)
    workflow.add_node("optimize", node_optimize)
    workflow.add_node("finalize", node_finalize)
    
    workflow.set_entry_point("load")
    
    workflow.add_edge("load", "convert")
    
    workflow.add_conditional_edges(
        "convert",
        should_fallback,
        {"fallback": "fallback", "validate": "validate"}
    )
    
    workflow.add_edge("fallback", "validate")
    
    workflow.add_conditional_edges(
        "validate",
        should_repair,
        {"repair": "repair", "split": "split"}
    )
    
    workflow.add_edge("repair", "split")
    workflow.add_edge("split", "optimize")
    workflow.add_edge("optimize", "finalize")
    workflow.add_edge("finalize", END)
    
    return workflow.compile()


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ë©”ì¸ í•¨ìˆ˜
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def process_document(
    filename: str,
    content: bytes,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    debug: bool = False,
    use_llm_metadata: bool = False,  # LLM ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ì‚¬ìš© ì—¬ë¶€
    use_clause_parsing: bool = True  # ì¡°í•­ ë²ˆí˜¸ ê¸°ë°˜ íŒŒì‹± ì‚¬ìš© ì—¬ë¶€
) -> dict:
    """ë¬¸ì„œ ì²˜ë¦¬ ë©”ì¸ í•¨ìˆ˜"""
    
    initial_state: PipelineState = {
        "filename": filename,
        "content": content,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "use_llm_metadata": use_llm_metadata,  # LLM ë©”íƒ€ë°ì´í„° ì¶”ì¶œ ì˜µì…˜
        "use_clause_parsing": use_clause_parsing,  # ì¡°í•­ ë²ˆí˜¸ ê¸°ë°˜ íŒŒì‹± ì˜µì…˜
        "file_type": "",
        "markdown": "",
        "metadata": {},
        "sections": [],
        "chunks": [],
        "quality_score": 0.0,
        "conversion_method": "",
        "errors": [],
        "warnings": [],
        "retry_count": 0,
        "success": False,
    }
    
    try:
        pipeline = build_pipeline()
        result = pipeline.invoke(initial_state)
        
        if debug:
            print(f"\n{'='*60}")
            print(f"ğŸ“Š LangGraph íŒŒì´í”„ë¼ì¸ ê²°ê³¼")
            print(f"{'='*60}")
            print(f"   íŒŒì¼: {filename}")
            print(f"   ë³€í™˜ ë°©ë²•: {result.get('conversion_method')}")
            print(f"   í’ˆì§ˆ ì ìˆ˜: {result.get('quality_score', 0):.0%}")
            print(f"   ì´ ì²­í¬: {len(result.get('chunks', []))}")
            if result.get('warnings'):
                print(f"   âš ï¸ ê²½ê³ : {result['warnings']}")
            if result.get('errors'):
                print(f"   âŒ ì—ëŸ¬: {result['errors']}")
        
        return result
        
    except ImportError:
        if debug:
            print("âš ï¸ LangGraph ì—†ìŒ, ì‹¬í”Œ íŒŒì´í”„ë¼ì¸ ì‚¬ìš©")
        return _simple_pipeline(initial_state, debug)


def _simple_pipeline(state: PipelineState, debug: bool = False) -> dict:
    """LangGraph ì—†ì„ ë•Œ ì‚¬ìš©í•˜ëŠ” ì‹¬í”Œ íŒŒì´í”„ë¼ì¸"""
    state = node_load(state)
    if state.get("errors"):
        return state
    
    state = node_convert(state)
    if not state.get("markdown"):
        state = node_convert_fallback(state)
    
    state = node_validate(state)
    
    if state.get("quality_score", 0) < 0.5:
        state = node_repair(state)
    
    state = node_split(state)
    state = node_optimize(state)
    state = node_finalize(state)
    
    if debug:
        print(f"\nğŸ“Š ì‹¬í”Œ íŒŒì´í”„ë¼ì¸ ê²°ê³¼: {len(state.get('chunks', []))} ì²­í¬")
    
    return state


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ê²°ê³¼ ë³€í™˜ í—¬í¼
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def state_to_chunks(state: dict) -> List[Chunk]:
    """ìƒíƒœë¥¼ Chunk ê°ì²´ ë¦¬ìŠ¤íŠ¸ë¡œ ë³€í™˜"""
    return [
        Chunk(
            text=c["text"],
            index=c["index"],
            metadata=c["metadata"]
        )
        for c in state.get("chunks", [])
    ]


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# í…ŒìŠ¤íŠ¸
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

if __name__ == "__main__":
    print("ğŸ”¥ document_pipeline v9.1 í…ŒìŠ¤íŠ¸")